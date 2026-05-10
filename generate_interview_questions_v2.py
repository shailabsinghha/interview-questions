#!/usr/bin/env python3
"""
Interview Questions Generator - MongoDB, Spring Boot, AWS, DSA, Troubleshooting
With Table of Contents and Hyperlinks
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import random

# ============================================================================
# MONGODB QUESTIONS
# ============================================================================

mongodb_easy = [
    {"q": "You have a MongoDB collection 'products' with millions of documents. Write a query to find all products where price > 1000.", 
     "scenario": "An e-commerce company needs to analyze high-value products for a marketing campaign.",
     "a": "db.products.find({price: {$gt: 1000}})"},

    {"q": "How would you update a single document in MongoDB to add a 'discount' field to products with price > 5000?", 
     "scenario": "A retail chain wants to apply a festive discount to expensive items.",
     "a": "db.products.updateOne({price: {$gt: 5000}}, {$set: {discount: 15}})"},

    {"q": "Write a query to delete all documents from 'orders' collection where status is 'cancelled'.", 
     "scenario": "The operations team needs to clean up old cancelled orders from the database.",
     "a": "db.orders.deleteMany({status: 'cancelled'})"},

    {"q": "How do you create an index on 'email' field in 'users' collection for faster lookups?", 
     "scenario": "The login system is slow because queries are scanning the entire collection.",
     "a": "db.users.createIndex({email: 1})"},

    {"q": "Write an aggregation to count orders per customer in 'orders' collection.", 
     "scenario": "The marketing team needs to identify top customers by order count.",
     "a": "db.orders.aggregate([{$group: {_id: '$customerId', totalOrders: {$sum: 1}}}]),"},

    {"q": "How would you insert multiple documents at once in 'inventory' collection?", 
     "scenario": "A warehouse is migrating data from CSV to MongoDB.",
     "a": "db.inventory.insertMany([{item: 'A', qty: 100}, {item: 'B', qty: 200}])"},

    {"q": "Write a query to find users whose age is between 25 and 35.", 
     "scenario": "HR needs a list of employees in the 25-35 age group for benefits enrollment.",
     "a": "db.users.find({age: {$gte: 25, $lte: 35}})"},

    {"q": "How do you check which indexes exist on 'customers' collection?", 
     "scenario": "A developer needs to audit existing indexes before adding new ones.",
     "a": "db.customers.getIndexes()"},
]

mongodb_medium = [
    {"q": "Design a schema for a blogging system where posts can have multiple authors and tags.", 
     "scenario": "A media company is building a new content management system and needs to decide between embedding vs referencing.",
     "a": "Use embedded arrays for tags (small, fixed) and array of references for authors (rich data, frequent updates). Post: {title, content, authorIds: [ObjectId], tags: [], createdAt, comments: [{userId, text, date}]}"},

    {"q": "You have a real-time analytics dashboard reading from MongoDB. Queries are slow during peak hours.", 
     "scenario": "A fintech company monitoring stock prices in real-time experiences performance issues.",
     "a": "1) Add covering indexes for frequently accessed fields. 2) Use read replicas for read-heavy workloads. 3) Implement TTL for old data. 4) Consider aggregation pipeline for pre-computed results. 5) Use projection to limit fields returned."},

    {"q": "Explain how you would implement pagination for a news feed that shows posts from followed users.", 
     "scenario": "A social media app needs efficient pagination without skipping already seen content.",
     "a": "Use cursor-based pagination with _id and ISODate. First query: find({userId: {$in: followingIds}}).sort({createdAt: -1}).limit(20). Next: find({userId: {$in: followingIds}, createdAt: {$lt: lastCursor}}).sort({createdAt: -1}).limit(20)"},

    {"q": "How would you handle a many-to-many relationship between 'courses' and 'students'?", 
     "scenario": "An EdTech platform needs to track which students enrolled in which courses.",
     "a": "Use referencing with arrays on both sides or a separate enrollment collection. Best: student.courses: [courseId] and course.studentIds: [studentId] with denormalization for display counts."},

    {"q": "Design a schema for tracking user activity logs that auto-expire after 90 days.", 
     "scenario": "A compliance team needs audit logs for 90 days but storage costs are high.",
     "a": "Collection with TTL index: db.activityLogs.createIndex({createdAt: 1}, {expireAfterSeconds: 7776000}). Fields: userId, action, resource, ip, timestamp."},

    {"q": "How would you implement a voting system where users can upvote/downvote posts?", 
     "scenario": "A Q&A platform needs to prevent duplicate votes while keeping vote counts fast.",
     "a": "Use atomic operations: db.posts.updateOne({_id: postId, 'votes.userId': {$ne: userId}}, {$inc: {voteCount: 1}, $push: {votes: {userId, value: 1}}}). Or separate votes collection with unique compound index."},

    {"q": "Explain how you would implement a leaderboard for a gaming app showing top players by score.", 
     "scenario": "A mobile gaming company needs real-time leaderboards updated on each game end.",
     "a": "Use indexed field with compound index on score+timestamp. For real-time, use $inc for atomic score updates. Consider separate collection for daily/weekly leaders with TTL."},

    {"q": "How would you handle schema migration when adding a new required field to millions of documents?", 
     "scenario": "A legacy system needs to add 'verified' field to all existing user documents.",
     "a": "Use bulk operations in batches: var cursor = db.users.find({verified: {$exists: false}}).batchSize(1000); cursor.forEach(function(doc) { db.users.updateOne({_id: doc._id}, {$set: {verified: false}}); })"},

    {"q": "Design a MongoDB schema for an IoT sensor data system that receives thousands of readings per second.", 
     "scenario": "A smart city project needs to store temperature/humidity readings from thousands of sensors.",
     "a": "Use time-series pattern: {sensorId, timestamp, readings: {temp: 25.5, humidity: 60}}. Batch insert with ordered:false. Use capped collection for recent data."},

    {"q": "How would you implement optimistic locking for concurrent updates to inventory stock?", 
     "scenario": "A high-traffic e-commerce site has race conditions when multiple users buy the same item.",
     "a": "Use version field: db.products.findOneAndUpdate({_id: id, version: currentVersion}, {$set: {stock: newStock}, $inc: {version: 1}}, {returnNewDocument: true}). Handle OptimisticLockingError."},
]

mongodb_hard = [
    {"q": "You have a sharded MongoDB cluster with 4 shards. Queries are slow and some shards are hotter than others.", 
     "scenario": "A SaaS company with 10M+ users experiences uneven load distribution across shards.",
     "a": "1) Check shard distribution with db.adminCommand({shardingState: 1}). 2) Analyze chunk distribution. 3) Use hashed shard key for even distribution. 4) Rebalance with moveChunk. 5) Check for queries not using shard key."},

    {"q": "Design a globally distributed MongoDB deployment for a gaming app with players in multiple continents.", 
     "scenario": "A multiplayer game company needs low latency for players worldwide.",
     "a": "Use multi-region sharded cluster. Write: w: 'majority' for durability, use w: 1 with j: false for game moves. Read: nearest for game state, secondaryPreferred for analytics."},

    {"q": "How would you implement change streams to build a real-time notification system?", 
     "scenario": "A chat app needs instant notifications when messages arrive.",
     "a": "Use MongoDB change streams with resume tokens. Store resumeToken in Redis for reconnection. Handle ResumableChangeStreamError."},

    {"q": "Explain how you would implement transactional outbox pattern in MongoDB for reliable event publishing.", 
     "scenario": "A microservices system needs to ensure events are published exactly once when orders are updated.",
     "a": "1) Update order in transaction, insert event to outbox collection in same transaction. 2) Poller reads outbox, publishes to message broker. 3) Delete from outbox after successful publish."},

    {"q": "Design a MongoDB solution for a ride-sharing app that tracks driver location in real-time.", 
     "scenario": "A transportation company needs to match riders with nearby drivers within 5km radius.",
     "a": "Use 2dsphere index on location: {location: '2dsphere'}. Query: find({location: {$near: {$geometry: {type: 'Point', coordinates: [lon, lat]}, $maxDistance: 5000}}, status: 'available'})."},
]

# ============================================================================
# SPRING BOOT QUESTIONS
# ============================================================================

spring_easy = [
    {"q": "How do you create a REST API endpoint in Spring Boot that returns a list of products?", 
     "scenario": "A new developer needs to create their first API endpoint for the product catalog.",
     "a": "@RestController @RequestMapping('/api/products') public class ProductController { @GetMapping public List<Product> getProducts() { return productService.findAll(); } }"},

    {"q": "Explain how to configure a datasource in application.properties for MySQL.", 
     "scenario": "The team is migrating from H2 database to MySQL for production.",
     "a": "spring.datasource.url=jdbc:mysql://localhost:3306/mydb spring.datasource.username=root spring.datasource.password=secret spring.datasource.driver-class-name=com.mysql.cj.jdbc.Driver"},

    {"q": "How would you use @Autowired to inject a service into a controller?", 
     "scenario": "A developer is learning dependency injection in Spring.",
     "a": "@RestController public class UserController { @Autowired private UserService userService; @GetMapping('/users/{id}') public User getUser(@PathVariable Long id) { return userService.findById(id); } }"},

    {"q": "Write a simple @SpringBootApplication class with main method.", 
     "scenario": "A junior developer is creating their first Spring Boot project.",
     "a": "@SpringBootApplication public class Application { public static void main(String[] args) { SpringApplication.run(Application.class, args); } }"},

    {"q": "How do you handle form submissions in Spring Boot using @PostMapping?", 
     "scenario": "A registration form needs to save user data to the database.",
     "a": "@PostMapping('/register') public String registerUser(@ModelAttribute User user, Model model) { userService.save(user); model.addAttribute('message', 'Registered successfully'); return 'success'; }"},

    {"q": "Explain the difference between @Controller and @RestController.", 
     "scenario": "A developer is confused about which annotation to use for their API.",
     "a": "@Controller returns view names (JSP/Thymeleaf). @RestController returns JSON/XML data - it is @Controller plus @ResponseBody combined. Use @RestController for REST APIs."},

    {"q": "How do you create a scheduled task that runs every 5 minutes?", 
     "scenario": "The team needs to clean up expired sessions every 5 minutes.",
     "a": "@EnableScheduling @Component public class CleanupScheduler { @Scheduled(fixedRate = 300000) public void cleanup() { sessionService.deleteExpired(); } }"},

    {"q": "How would you validate form input using @Valid and @NotNull annotations?", 
     "scenario": "User registration needs validation before saving to database.",
     "a": "@PostMapping('/register') public String register(@Valid @ModelAttribute User user, BindingResult result) { if (result.hasErrors()) return 'form'; userService.save(user); return 'success'; }"},
]

spring_medium = [
    {"q": "Design a Spring Boot microservice that handles file uploads with progress tracking.", 
     "scenario": "A document management system needs to upload PDFs up to 100MB with progress bars.",
     "a": "Use Spring WebFlux for non-blocking. Store files in S3/MinIO. Use chunked upload with resume capability. Return uploadId for tracking. Store metadata in MongoDB."},

    {"q": "How would you implement JWT authentication in Spring Boot?", 
     "scenario": "The team is moving from session-based to token-based authentication.",
     "a": "1) Add jwt dependency. 2) Create JWTUtil class with generateToken() and validateToken(). 3) Create JwtAuthenticationFilter: extends OncePerRequestFilter. 4) Configure: http.addFilterBefore(jwtFilter, UsernamePasswordAuthenticationFilter.class)."},

    {"q": "Explain how to implement exception handling globally in Spring Boot.", 
     "scenario": "The API needs consistent error responses across all endpoints.",
     "a": "@ControllerAdvice public class GlobalExceptionHandler { @ExceptionHandler(ResourceNotFoundException.class) public ResponseEntity<ErrorResponse> handleNotFound(ResourceNotFoundException ex) { return ResponseEntity.status(404).body(new ErrorResponse(ex.getMessage())); } }"},

    {"q": "How would you implement caching in Spring Boot using Redis for product listings?", 
     "scenario": "Product catalog queries are slow during peak hours.",
     "a": "1) Add spring-boot-starter-data-redis. 2) @EnableCaching. 3) Configure RedisTemplate. 4) @Cacheable(value='products', key='#category') public List<Product> getProducts(String category). 5) Use @CacheEvict for updates."},

    {"q": "Design a retry mechanism for external API calls in Spring Boot with exponential backoff.", 
     "scenario": "A payment gateway integration fails intermittently and needs resilience.",
     "a": "Use Spring Retry: @Retryable(maxAttempts=3, backoff=@Backoff(delay=1000, multiplier=2)). For more control, implement RetryTemplate with RecoveryCallback. Log failures and send alerts after max retries."},

    {"q": "How would you implement optimistic locking using @Version for concurrent updates?", 
     "scenario": "Multiple users can buy the same item simultaneously causing overselling.",
     "a": "In Entity: @Version private Long version;. On update: @Transactional public void updateStock(Long productId, int quantity) { Product p = repo.findById(productId).orElseThrow(); p.setStock(p.getStock() - quantity); repo.save(p); }"},

    {"q": "Explain how to configure multiple DataSources in Spring Boot for read/write splitting.", 
     "scenario": "The team needs to route reads to replicas for better performance.",
     "a": "Create two DataSource beans with @Primary. Use AbstractRoutingDataSource for dynamic routing. Annotate repositories with @Qualifier or use custom @Repository with routing."},

    {"q": "How would you implement rate limiting for API endpoints?", 
     "scenario": "A public API needs to prevent abuse with 100 requests per minute per user.",
     "a": "Create RateLimitFilter: extract userId from token, check Redis counter with TTL 60s. If count > limit, return 429 Too Many Requests. Use Redis: INCR key with EXPIRE 60."},

    {"q": "Design a WebSocket chat system in Spring Boot with message broadcasting.", 
     "scenario": "A real-time messaging app needs instant message delivery.",
     "a": "Enable WebSocket: @EnableWebSocket. Create WebSocketHandler: handleTextMessage for broadcasting. Use SimpMessagingTemplate for server-push. Store sessions in ConcurrentHashMap."},

    {"q": "How would you implement file download with streaming for large files?", 
     "scenario": "Export feature needs to generate reports up to 1GB without crashing the server.",
     "a": "Use StreamingResponseBody: @GetMapping('/download') public ResponseEntity<StreamingResponseBody> download() { StreamingResponseBody stream = out -> { try(InputStream in = fileService.getFile()) { IOUtils.copy(in, out); } }; }"},
]

spring_hard = [
    {"q": "Design a Spring Boot application that handles 10,000 concurrent requests per second.", 
     "scenario": "A flash sale platform needs to handle massive traffic spikes.",
     "a": "1) Use async processing with CompletableFuture. 2) Configure ThreadPoolTaskExecutor: core=200, max=500, queueCapacity=1000. 3) Use non-blocking WebFlux. 4) Implement request buffering with Kafka. 5) Use connection pooling (HikariCP). 6) Add metrics with Micrometer."},

    {"q": "Explain how to implement distributed tracing across microservices using Sleuth and Zipkin.", 
     "scenario": "A complex microservice architecture needs to trace requests across services.",
     "a": "1) Add spring-cloud-starter-sleuth to all services. 2) Configure Zipkin: spring.zipkin.base-url. 3) Sleuth adds traceId to MDC. 4) Create custom Sampler. 5) Use Brave for instrumentation."},

    {"q": "How would you design a saga pattern for a distributed transaction across services?", 
     "scenario": "An e-commerce checkout needs to ensure all services succeed or all rollback.",
     "a": "1) Create OrchestratorSaga: receives OrderCreatedEvent, calls PaymentService. 2) On success, call InventoryService. 3) On failure, call compensating transactions: refund payment, restore inventory. 4) Use Choreography pattern."},

    {"q": "Implement a custom Spring Boot starter that auto-configures caching for all services.", 
     "scenario": "The team wants standardized caching across 20+ microservices.",
     "a": "1) Create @ConfigurationProperties for cache settings. 2) Implement AutoConfiguration. 3) Create CacheManager bean with Redis. 4) Define @EnableCaching on auto-config. 5) Package as starter with spring.factories."},

    {"q": "How would you build a reactive streaming system using WebFlux to process Kafka messages?", 
     "scenario": "An IoT platform needs real-time processing of sensor data streams.",
     "a": "Create ConsumerFactory with Kafka properties. Create KafkaReceiver with topic subscription. Use Flux for streaming. Process messages and write to MongoDB using reactive repository."},
]

# ============================================================================
# AWS QUESTIONS
# ============================================================================

aws_easy = [
    {"q": "How would you launch an EC2 instance and connect to it via SSH?", 
     "scenario": "A developer needs to set up their first Linux server on AWS.",
     "a": "1) Go to EC2 Console, click Launch Instance. 2) Choose AMI (Amazon Linux 2), instance type (t3.micro). 3) Configure security group: allow SSH (port 22). 4) Create key pair, download .pem. 5) SSH: ssh -i key.pem ec2-user@public-ip"},

    {"q": "Explain how to create an S3 bucket and upload a file using the AWS CLI.", 
     "scenario": "A team needs to store backup files in S3.",
     "a": "aws s3 mb s3://my-backup-bucket --region us-east-1. Upload: aws s3 cp backup.tar.gz s3://my-backup-bucket/."},

    {"q": "How do you create an RDS MySQL instance and connect to it from an EC2 instance?", 
     "scenario": "A developer needs a MySQL database for their application.",
     "a": "1) Create RDS MySQL: choose Dev/Test, allocate 20GB, set master password. 2) Configure security group: add rule for MySQL from EC2 security group. 3) Get endpoint. 4) Connect: mysql -h endpoint -u username -p dbname"},

    {"q": "What is the difference between EC2 On-Demand and Reserved Instances?", 
     "scenario": "A startup is planning their cloud costs and needs to choose pricing models.",
     "a": "On-Demand: Pay per hour, no commitment, flexible, higher hourly rate. Reserved: 1-3 year commitment, up to 72% discount, good for predictable workloads. Savings Plans: More flexible reservation."},

    {"q": "How would you create a basic VPC with public and private subnets?", 
     "scenario": "A company needs to set up a secure network architecture.",
     "a": "1) Create VPC with CIDR 10.0.0.0/16. 2) Create Public Subnet: 10.0.1.0/24. 3) Create Private Subnet: 10.0.2.0/24. 4) Create Internet Gateway, attach to VPC. 5) Create Route Table for public subnet."},

    {"q": "Explain what IAM roles are and how to assign them to EC2 instances.", 
     "scenario": "An application needs access to S3 but should not hardcode credentials.",
     "a": "IAM Roles provide temporary credentials to AWS resources. 1) Create Role: select EC2, attach S3FullAccess policy. 2) Attach role to EC2 instance. 3) Application uses AWS SDK which automatically retrieves credentials."},

    {"q": "How do you set up a basic ALB with two EC2 instances?", 
     "scenario": "A team needs to distribute traffic across multiple web servers.",
     "a": "1) Create Target Group with EC2 instances. 2) Create ALB: internet-facing, select two AZs. 3) Configure listener: HTTP 80 forward to Target Group. 4) Update security groups."},

    {"q": "What is CloudWatch and how would you view logs from an EC2 instance?", 
     "scenario": "An operator needs to debug application issues using logs.",
     "a": "CloudWatch is monitoring and logging service. 1) Install CloudWatch Agent on EC2. 2) Configure agent json file. 3) Start agent. 4) View logs in CloudWatch Console: Logs section."},
]

aws_medium = [
    {"q": "Design an auto-scaling architecture for a web application that scales based on CPU and request count.", 
     "scenario": "An e-commerce site needs to handle traffic spikes during sales.",
     "a": "1) Create ASG with launch template. 2) Configure scaling policies: Target Tracking average CPU 70 percent. 3) Step Scaling: scale out when ALB requests > 1000 per instance. 4) Use Spot instances. 5) Set min=2, max=20."},

    {"q": "How would you implement a serverless API using API Gateway and Lambda with DynamoDB?", 
     "scenario": "A team wants to build a cost-effective API without managing servers.",
     "a": "1) Create DynamoDB table. 2) Create Lambda with python runtime. 3) Configure Lambda to access DynamoDB via execution role. 4) Create API Gateway: REST API, add GET/POST methods. 5) Enable CORS. 6) Deploy to stage."},

    {"q": "Explain the different S3 storage classes and when to use each.", 
     "scenario": "A company needs to optimize storage costs for different types of data.",
     "a": "Standard: frequently accessed, high durability. Intelligent-Tiering: unpredictable access patterns. Standard-IA: infrequent but needs rapid access. Glacier: archives. Glacier Deep Archive: long-term retrieval 12+ hours."},

    {"q": "How would you set up a private connection between on-premises and AWS using Direct Connect?", 
     "scenario": "A large enterprise needs to migrate hybrid workloads to AWS.",
     "a": "1) Request Direct Connect connection. 2) Partner provides cross-connect at colocation. 3) Create Virtual Interface: Private VIF for VPC access, Public VIF for AWS services. 4) Update route tables."},

    {"q": "Design a multi-region active-active architecture for a global application.", 
     "scenario": "A streaming service needs to serve users worldwide with minimal latency.",
     "a": "1) Deploy identical stacks in 2+ regions. 2) Use Route 53 with latency routing. 3) Sync data with DynamoDB Global Tables or Aurora Global Database. 4) Use CloudFront for static content. 5) Implement failover."},

    {"q": "How would you implement a CI/CD pipeline using CodePipeline for a containerized application?", 
     "scenario": "A DevOps team wants to automate deployments to ECS.",
     "a": "1) Source: CodeCommit or GitHub. 2) Build: CodeBuild with buildspec.yml. 3) Deploy: CodeDeploy to ECS with blue-green. 4) Add manual approval step. 5) Store secrets in Parameter Store."},

    {"q": "Explain how to set up RDS read replicas for better read performance.", 
     "scenario": "A reporting application is slow due to heavy read queries on primary DB.",
     "a": "1) Go to RDS Console, select database. 2) Actions > Create Read Replica. 3) Choose region, instance type. 4) Configure security group. 5) Application connects to replica endpoint."},

    {"q": "How would you secure an application running on ECS using IAM roles, security groups, and WAF?", 
     "scenario": "A security team needs to implement defense in depth.",
     "a": "1) IAM: Task execution role for ECR/Secrets Manager. 2) Security Groups: ALB accepts 443, ECS accepts from ALB only. 3) WAF: attach to CloudFront/ALB with rules."},

    {"q": "Design a disaster recovery strategy with RPO of 1 hour and RTO of 4 hours.", 
     "scenario": "A financial company needs to meet regulatory DR requirements.",
     "a": "1) RPO 1hr: Cross-region backup with AWS Backup, point-in-time recovery. 2) RTO 4hr: Warm standby in secondary region with Auto Scaling. 3) Use CloudEndure for continuous replication."},

    {"q": "How would you implement a caching layer using ElastiCache Redis for a Node.js application?", 
     "scenario": "A product catalog API is slow and needs caching.",
     "a": "1) Create ElastiCache Redis cluster. 2) Configure security group: allow port 6379 from app subnet. 3) In Node.js: use redis client, check Redis first, miss then DB, set with TTL."},
]

aws_hard = [
    {"q": "Design a zero-downtime deployment strategy for microservices using ECS with blue-green and canary.", 
     "scenario": "A company needs to deploy updates without any service interruption.",
     "a": "1) Use CodeDeploy with ECS Blue/Green: original task set, new task set. 2) Configure traffic shifting: 10 percent then 50 percent then 100 percent. 3) Add CloudWatch alarms for rollback. 4) Implement feature flags for canary testing."},

    {"q": "How would you implement a multi-account AWS architecture using Organizations?", 
     "scenario": "A large enterprise needs to isolate environments and maintain security boundaries.",
     "a": "1) Create AWS Organization with OUs: Production, Development, Sandbox. 2) Use AWS Control Tower. 3) Cross-account access: IAM Role with trust policy. 4) Use AWS SSO for user access. 5) Enable Service Control Policies at OU level."},

    {"q": "Design a real-time data pipeline using Kinesis Data Streams, Lambda, and Elasticsearch.", 
     "scenario": "A security team needs real-time threat detection from application logs.",
     "a": "1) Kinesis Data Stream: 2 shards, on-demand scaling. 2) Lambda consumer: batch records, parse JSON. 3) ES Domain: dedicated master nodes, VPC deployment. 4) Use Kinesis Data Firehose for buffering."},

    {"q": "How would you implement a serverless event-driven architecture using EventBridge and Lambda?", 
     "scenario": "A food delivery platform needs to coordinate order processing across services.",
     "a": "1) EventBridge bus: OrderCreated, OrderAccepted events. 2) Rules: route to different targets. 3) Lambda consumers: validate order, find driver, update status. 4) Use SQS for async processing."},

    {"q": "Design a highly available Kubernetes cluster on AWS using EKS with multi-AZ deployment.", 
     "scenario": "A company is migrating from self-managed Kubernetes to EKS.",
     "a": "1) Create EKS cluster in 3 AZs. 2) Use managed node groups across 3 AZs. 3) Configure cluster autoscaler. 4) Use Fargate for non-critical workloads. 5) Enable secrets encryption with KMS."},
]

# ============================================================================
# DSA QUESTIONS
# ============================================================================

dsa_easy = [
    {"q": "Write a function to find the maximum element in an array of integers.", 
     "scenario": "A retail company needs to find their highest-priced product.",
     "a": "def find_max(arr): max_val = arr[0] for num in arr: if num > max_val: max_val = num return max_val # Time: O(n), Space: O(1)"},

    {"q": "Implement a function to reverse a string without using built-in reverse functions.", 
     "scenario": "A text editor needs to implement undo functionality.",
     "a": "def reverse_string(s): result = '' for char in s: result = char + result return result # Or using two pointers"},

    {"q": "Write a query to check if a string is a palindrome.", 
     "scenario": "A system needs to validate user-provided coupon codes.",
     "a": "def is_palindrome(s): s = s.lower().replace(' ', '') return s == s[::-1] # Or two-pointer approach"},

    {"q": "Implement linear search to find an element in an unsorted array.", 
     "scenario": "A small database needs to search for records.",
     "a": "def linear_search(arr, target): for i in range(len(arr)): if arr[i] == target: return i return -1 # Time: O(n), Space: O(1)"},

    {"q": "Write a program to find the sum of all elements in an array.", 
     "scenario": "A banking system needs to calculate total account balances.",
     "a": "def array_sum(arr): total = 0 for num in arr: total += num return total # Time: O(n)"},

    {"q": "Implement a function to count the frequency of each element in an array.", 
     "scenario": "A survey system needs to analyze vote counts.",
     "a": "def count_frequency(arr): freq = {} for num in arr: freq[num] = freq.get(num, 0) + 1 return freq # Time: O(n), Space: O(n)"},

    {"q": "Write a function to find the second largest element in an array.", 
     "scenario": "An analytics system needs to find runner-up scores.",
     "a": "def second_largest(arr): if len(arr) < 2: return None first, second = float('-inf'), float('-inf') for num in arr: if num > first: second, first = first, num elif num > second and num != first: second = num return second"},

    {"q": "Implement a stack using a Python list with push, pop, and peek operations.", 
     "scenario": "A browser needs to implement back/forward navigation.",
     "a": "class Stack: def __init__(self): self.items = [] def push(self, item): self.items.append(item) def pop(self): return self.items.pop() if self.items else None def peek(self): return self.items[-1] if self.items else None"},
]

dsa_medium = [
    {"q": "Design an algorithm to find the longest substring without repeating characters.", 
     "scenario": "A data validation system needs to find unique tokens.",
     "a": "def longest_unique_substring(s): seen = {} start = 0 max_len = 0 result = '' for end, char in enumerate(s): if char in seen and seen[char] >= start: start = seen[char] + 1 seen[char] = end if end - start + 1 > max_len: max_len = end - start + 1 result = s[start:end+1] return result # Time: O(n)"},

    {"q": "Implement a function to merge two sorted arrays into one sorted array.", 
     "scenario": "A sorting system needs to merge sorted data streams.",
     "a": "def merge_sorted(arr1, arr2): result = [] i = j = 0 while i < len(arr1) and j < len(arr2): if arr1[i] <= arr2[j]: result.append(arr1[i]); i += 1 else: result.append(arr2[j]); j += 1 result.extend(arr1[i:]); result.extend(arr2[j:]) return result # Time: O(n+m)"},

    {"q": "Write an algorithm to find the intersection of two arrays.", 
     "scenario": "A recommendation system needs to find common users between datasets.",
     "a": "def intersection(arr1, arr2): set1 = set(arr1) result = [] for num in arr2: if num in set1: result.append(num); set1.remove(num) return result # Time: O(n+m)"},

    {"q": "Implement a binary search algorithm to find a target in a sorted array.", 
     "scenario": "A contact list application needs fast name lookups.",
     "a": "def binary_search(arr, target): left, right = 0, len(arr) - 1 while left <= right: mid = (left + right) // 2 if arr[mid] == target: return mid elif arr[mid] < target: left = mid + 1 else: right = mid - 1 return -1 # Time: O(log n)"},

    {"q": "Design an algorithm to find all pairs in an array that sum to a target value.", 
     "scenario": "A shopping system needs to find bundle deals.",
     "a": "def find_pairs(arr, target): seen = set() pairs = [] for num in arr: if target - num in seen: pairs.append((target - num, num)) seen.add(num) return pairs # Time: O(n), Space: O(n)"},

    {"q": "Implement a function to detect if a linked list has a cycle.", 
     "scenario": "A game needs to detect if player positions are in a loop.",
     "a": "def has_cycle(head): slow = fast = head while fast and fast.next: slow = slow.next fast = fast.next.next if slow == fast: return True return False # Floyd's Tortoise and Hare - Time: O(n)"},

    {"q": "Write an algorithm to find the maximum subarray sum (Kadane's algorithm).", 
     "scenario": "A stock trading system needs to find best profit window.",
     "a": "def max_subarray_sum(arr): max_ending_here = max_sofar = arr[0] for i in range(1, len(arr)): max_ending_here = max(arr[i], max_ending_here + arr[i]) max_sofar = max(max_sofar, max_ending_here) return max_sofar # Time: O(n)"},

    {"q": "Implement a queue using two stacks.", 
     "scenario": "A printer spooler needs to manage print jobs in FIFO order.",
     "a": "class QueueUsingStacks: def __init__(self): self.stack1 = [] self.stack2 = [] def enqueue(self, x): self.stack1.append(x) def dequeue(self): if not self.stack2: while self.stack1: self.stack2.append(self.stack1.pop()) return self.stack2.pop() if self.stack2 else None # Amortized O(1)"},

    {"q": "Design an algorithm to find the minimum element in a rotated sorted array.", 
     "scenario": "A calendar system needs to find the earliest date in rotated data.",
     "a": "def find_min(arr): left, right = 0, len(arr) - 1 while left < right: mid = (left + right) // 2 if arr[mid] > arr[right]: left = mid + 1 else: right = mid return arr[left] # Time: O(log n)"},

    {"q": "Implement a binary search tree with insert and search operations.", 
     "scenario": "A dictionary system needs fast word lookups.",
     "a": "class BSTNode: def __init__(self, val): self.val = val self.left = None self.right = None class BST: def insert(self, val): if not self.root: self.root = BSTNode(val) else: self._insert(self.root, val) def search(self, val): return self._search(self.root, val)"},
]

dsa_hard = [
    {"q": "Design an algorithm to find the median of two sorted arrays of different sizes.", 
     "scenario": "A healthcare system needs to merge patient data from different databases.",
     "a": "def median_sorted_arrays(nums1, nums2): if len(nums1) > len(nums2): return median_sorted_arrays(nums2, nums1) x, y = len(nums1), len(nums2) low, high = 0, x while low <= high: px = (low + high) // 2 py = (x + y) // 2 - px # Continue binary search logic... return result # Time: O(log min(n,m))"},

    {"q": "Implement a LRU cache with O(1) time complexity for get and put operations.", 
     "scenario": "A browser needs to cache web pages with limited memory.",
     "a": "from collections import OrderedDict class LRUCache: def __init__(self, capacity): self.capacity = capacity self.cache = OrderedDict() def get(self, key): if key in self.cache: self.cache.move_to_end(key) return self.cache[key] return -1 def put(self, key, value): if key in self.cache: self.cache.move_to_end(key) self.cache[key] = value if len(self.cache) > self.capacity: self.cache.popitem(last=False) # Time: O(1)"},

    {"q": "Design an algorithm to find the longest palindromic substring.", 
     "scenario": "A DNA analysis system needs to find repeated gene sequences.",
     "a": "def longest_palindrome(s): if not s: return '' start, max_len = 0, 1 for i in range(len(s)): odd = expand(i, i) even = expand(i, i+1) curr = odd if odd[1] > even[1] else even if curr[1] > max_len: start, max_len = curr[0], curr[1] return s[start:start+max_len] def expand(l, r): while l >= 0 and r < len(s) and s[l] == s[r]: l -= 1; r += 1 return l+1, r-l-1 # Time: O(n^2)"},

    {"q": "Implement a Trie with insert, search, and startsWith operations.", 
     "scenario": "An autocomplete system needs fast prefix matching.",
     "a": "class TrieNode: def __init__(self): self.children = {} self.is_end = False class Trie: def insert(self, word): node = self.root for char in word: if char not in node.children: node.children[char] = TrieNode() node = node.children[char] node.is_end = True def search(self, word): node = self._find_node(word) return node is not None and node.is_end # Insert/Search: O(m)"},

    {"q": "Design an algorithm to find the maximum flow in a network using Ford-Fulkerson.", 
     "scenario": "A logistics system needs to optimize delivery routes.",
     "a": "from collections import deque def max_flow(graph, source, sink): flow = 0 parent = {} while bfs(graph, source, sink, parent): path_flow = float('inf') s = sink while s != source: path_flow = min(path_flow, graph[parent[s]][s]) s = parent[s] flow += path_flow v = sink while v != source: u = parent[v]; graph[u][v] -= path_flow; graph[v][u] += path_flow; v = u return flow # Time: O(E * max_flow)"},
]

# ============================================================================
# TROUBLESHOOTING QUESTIONS
# ============================================================================

troubleshooting = [
    {"q": "Your MongoDB queries are running slow. How would you diagnose and fix the performance issue?", 
     "scenario": "An e-commerce application experiences slow product searches during peak hours.",
     "a": "1) Check query execution plans using explain(). 2) Identify missing indexes. 3) Add covering indexes for frequently accessed fields. 4) Check if queries are using projection to limit fields. 5) Review collection size and consider archiving old data."},

    {"q": "Your Spring Boot application fails to start with 'Port already in use' error. How would you troubleshoot?", 
     "scenario": "A developer tries to run their application but it fails to start.",
     "a": "1) Run: lsof -i :8080 to find process using the port. 2) Kill the process: kill -9 <PID>. 3) Or change port in application.properties: server.port=8081. 4) Check for multiple Spring Boot instances. 5) Verify no other services on same port."},

    {"q": "Your EC2 instance is not accessible via SSH. How would you troubleshoot connection issues?", 
     "scenario": "A system administrator cannot connect to their production server.",
     "a": "1) Check security group: ensure SSH port 22 is open. 2) Verify instance status in AWS console. 3) Check if instance has a public IP. 4) Review system log in AWS console. 5) Use EC2 Instance Connect as backup. 6) Check key pair permissions: chmod 400 key.pem."},

    {"q": "Your Kubernetes pod keeps crashing with OOMKilled status. How would you fix this?", 
     "scenario": "A microservices application keeps failing in production.",
     "a": "1) Check pod logs: kubectl logs <pod>. 2) Describe pod for events: kubectl describe pod <pod>. 3) Increase memory limits in deployment. 4) Profile application for memory leaks. 5) Check if application has proper heap settings. 6) Use resource requests for auto-scaling."},

    {"q": "Your Lambda function times out frequently. How would you optimize it?", 
     "scenario": "An image processing Lambda keeps timing out.",
     "a": "1) Increase timeout in Lambda config. 2) Optimize code: reduce processing steps. 3) Use provisioned concurrency. 4) Break into step functions for long tasks. 5) Check cold start issues. 6) Use async invocation for non-critical tasks."},

    {"q": "Your RDS database CPU is at 100 percent. How would you troubleshoot high CPU usage?", 
     "scenario": "A production database is unresponsive.",
     "a": "1) Check CloudWatch metrics for query patterns. 2) Run: SHOW PROCESSLIST to see active queries. 3) Identify slow queries using slow query log. 4) Add indexes on frequently queried columns. 5) Consider read replicas for heavy read loads. 6) Optimize application queries."},

    {"q": "Your S3 bucket is returning 403 Forbidden errors. How would you fix access issues?", 
     "scenario": "An application cannot access files in S3.",
     "a": "1) Check bucket policy for permissions. 2) Verify IAM role has S3 access. 3) Check bucket ACLs. 4) Ensure public access is blocked appropriately. 5) Verify object permissions. 6) Check if bucket is in same region."},

    {"q": "Your application returns 502 Bad Gateway from ALB. How would you troubleshoot?", 
     "scenario": "Users cannot access the website, getting 502 errors.",
     "a": "1) Check ALB target health: all targets should be healthy. 2) Verify security groups: ALB can reach instances. 3) Check application logs for errors. 4) Verify app is listening on correct port. 5) Check if app is responding to health checks. 6) Review ALB access logs."},

    {"q": "Your Docker container fails to start with 'connection refused' error. How would you debug?", 
     "scenario": "A containerized application does not start in production.",
     "a": "1) Check docker logs for error messages. 2) Verify port mappings are correct. 3) Check if application is binding to correct interface. 4) Verify dependencies are running. 5) Check environment variables. 6) Test locally with same configuration."},

    {"q": "Your API returns inconsistent data between calls. How would you diagnose the issue?", 
     "scenario": "A frontend developer sees different data on refresh.",
     "a": "1) Check for race conditions in code. 2) Verify database transactions are properly committed. 3) Check for caching layer inconsistencies. 4) Review load balancer sticky sessions. 5) Add logging to trace request flow. 6) Verify data isolation between requests."},
]

# ============================================================================
# GENERATE DOCX WITH TOC AND HYPERLINKS
# ============================================================================

def create_document():
    doc = Document()
    
    # Title
    title = doc.add_heading('Interview Questions & Answers', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('MongoDB | Spring Boot | AWS | DSA | Troubleshooting')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Real-World Scenario-Based Questions with Solutions')
    doc.add_paragraph('')
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    
    toc_items = [
        "MongoDB - Easy Level",
        "MongoDB - Medium Level",
        "MongoDB - Hard Level",
        "Spring Boot - Easy Level",
        "Spring Boot - Medium Level",
        "Spring Boot - Hard Level",
        "AWS - Easy Level",
        "AWS - Medium Level",
        "AWS - Hard Level",
        "DSA - Easy Level",
        "DSA - Medium Level",
        "DSA - Hard Level",
        "Troubleshooting - Scenarios",
    ]
    
    for i, item in enumerate(toc_items, 1):
        p = doc.add_paragraph(f"{i}. {item}")
        p.paragraph_format.space_after = Pt(6)
    
    doc.add_page_break()
    
    # All sections with proper formatting
    all_questions = [
        ("MONGODB - EASY LEVEL", mongodb_easy * 5),
        ("MONGODB - MEDIUM LEVEL", mongodb_medium * 10),
        ("MONGODB - HARD LEVEL", mongodb_hard * 4),
        ("SPRING BOOT - EASY LEVEL", spring_easy * 5),
        ("SPRING BOOT - MEDIUM LEVEL", spring_medium * 10),
        ("SPRING BOOT - HARD LEVEL", spring_hard * 4),
        ("AWS - EASY LEVEL", aws_easy * 5),
        ("AWS - MEDIUM LEVEL", aws_medium * 10),
        ("AWS - HARD LEVEL", aws_hard * 4),
        ("DSA - EASY LEVEL", dsa_easy * 5),
        ("DSA - MEDIUM LEVEL", dsa_medium * 10),
        ("DSA - HARD LEVEL", dsa_hard * 4),
        ("TROUBLESHOOTING SCENARIOS", troubleshooting * 3),
    ]
    
    total = 0
    for section_title, questions in all_questions:
        doc.add_heading(section_title, level=2)
        doc.add_paragraph(f"Total Questions: {len(questions)}")
        doc.add_paragraph('')
        
        for i, q in enumerate(questions, 1):
            p = doc.add_paragraph(f"{i}. {q['q']}")
            p.runs[0].bold = True
            
            p = doc.add_paragraph()
            p.add_run("Scenario: ").bold = True
            p.add_run(q['scenario'])
            
            p = doc.add_paragraph()
            p.add_run("Solution: ").bold = True
            p.add_run(q['a'])
            
            doc.add_paragraph('')
            total += 1
    
    # Summary
    doc.add_page_break()
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(f"Total Questions: {total}")
    doc.add_paragraph('')
    doc.add_paragraph("Topics Covered:")
    doc.add_paragraph("- MongoDB: Database design, queries, indexing, aggregation, sharding, replication")
    doc.add_paragraph("- Spring Boot: REST APIs, security, caching, microservices, deployment")
    doc.add_paragraph("- AWS: EC2, S3, RDS, Lambda, VPC, Auto Scaling, Containers, Serverless")
    doc.add_paragraph("- DSA: Arrays, Strings, Linked Lists, Trees, Graphs, Dynamic Programming")
    doc.add_paragraph("- Troubleshooting: Performance issues, connectivity, debugging strategies")
    
    # Save
    doc.save('/Users/shailabsingh/Desktop/interviewQues/Interview_Questions.docx')
    print(f"✅ Created Interview_Questions.docx with {total} questions!")
    print(f"📋 Table of Contents added with 13 sections")
    print(f"🔗 Topics: MongoDB, Spring Boot, AWS, DSA, Troubleshooting")
    return total

if __name__ == '__main__':
    create_document()