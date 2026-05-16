#!/usr/bin/env python3
import http.server
import socketserver
import webbrowser
import os
import json
import io
import random
import urllib.parse
import urllib.request
import urllib.error
import subprocess
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

PORT = 8000
os.chdir('/Users/shailabsingh/Desktop/interviewQues')

class CustomHandler(http.server.SimpleHTTPRequestHandler):
    def do_GET(self):
        if self.path == '/api/questions':
            try:
                with open('questions.json', 'r', encoding='utf-8') as f:
                    questions = json.load(f)
                random.shuffle(questions)
                response = json.dumps(questions).encode('utf-8')
                self.send_response(200)
                self.send_header('Content-Type', 'application/json')
                self.send_header('Access-Control-Allow-Origin', '*')
                self.send_header('Content-Length', str(len(response)))
                self.end_headers()
                self.wfile.write(response)
            except Exception as e:
                self.send_response(500)
                self.send_header('Content-Type', 'text/plain')
                self.end_headers()
                self.wfile.write(str(e).encode())
        else:
            super().do_GET()

    def do_POST(self):
        content_length = int(self.headers.get('Content-Length', 0))
        post_data = self.rfile.read(content_length).decode('utf-8')

        if self.path == '/api/ai/generate':
            try:
                body = json.loads(post_data)
            except json.JSONDecodeError:
                self.send_response(400)
                self.send_header('Content-Type', 'application/json')
                self.send_header('Access-Control-Allow-Origin', '*')
                self.end_headers()
                self.wfile.write(json.dumps({'error': 'Invalid JSON'}).encode())
                return

            try:
                prompt = body.get('prompt', '')
                model = body.get('model', 'meta-llama/llama-3.3-70b-instruct:free')
                api_key = body.get('apiKey', '')

                result = None
                err_msg = None

                # 1) Try OpenRouter if an API key was provided
                if api_key:
                    openrouter_model = model.replace('opencode/', '')
                    payload = json.dumps({
                        'model': openrouter_model,
                        'messages': [{'role': 'user', 'content': prompt}],
                        'max_tokens': 4096,
                    }).encode()
                    req = urllib.request.Request(
                        'https://openrouter.ai/api/v1/chat/completions',
                        data=payload,
                        headers={
                            'Content-Type': 'application/json',
                            'Authorization': f'Bearer {api_key}',
                        },
                        method='POST',
                    )
                    try:
                        with urllib.request.urlopen(req, timeout=120) as resp:
                            or_data = json.loads(resp.read().decode())
                            result = or_data.get('choices', [{}])[0].get('message', {}).get('content', '')
                    except urllib.error.HTTPError as e:
                        err_msg = f'OpenRouter API error ({e.code}): {e.read().decode()}'
                    except Exception as e:
                        err_msg = f'OpenRouter error: {str(e)}'

                # 2) Fallback: opencode subprocess (local only)
                if result is None:
                    OPENCODE_BIN = '/opt/homebrew/lib/node_modules/opencode-ai/bin/opencode.exe'
                    oc_model = model if model.startswith('opencode/') else 'opencode/deepseek-v4-flash-free'
                    cmd = [OPENCODE_BIN, 'run', '--model', oc_model, '--format', 'json']
                    proc = subprocess.run(cmd, input=prompt, capture_output=True, text=True, timeout=120, env={**os.environ})
                    if proc.returncode == 0:
                        text_parts = []
                        for line in proc.stdout.strip().split('\n'):
                            line = line.strip()
                            if not line:
                                continue
                            try:
                                event = json.loads(line)
                                if event.get('type') == 'text':
                                    text_parts.append(event.get('part', {}).get('text', ''))
                            except json.JSONDecodeError:
                                continue
                        result = ''.join(text_parts) or None
                    else:
                        err_msg = (err_msg or '') + ('; ' if err_msg else '') + f'opencode failed: {proc.stderr[:200]}'

                if result is None:
                    raise Exception(err_msg or 'All AI backends failed')

                self.send_response(200)
                self.send_header('Content-Type', 'application/json')
                self.send_header('Access-Control-Allow-Origin', '*')
                self.end_headers()
                self.wfile.write(json.dumps({'text': result}).encode())

            except Exception as e:
                self.send_response(500)
                self.send_header('Content-Type', 'application/json')
                self.send_header('Access-Control-Allow-Origin', '*')
                self.end_headers()
                self.wfile.write(json.dumps({'error': str(e)}).encode())

        elif self.path == '/download':
            
            try:
                params = urllib.parse.parse_qs(post_data)
                question_ids_str = params.get('questions', [''])[0]
                question_ids = [int(x) for x in question_ids_str.split(',') if x]
                candidate_name = params.get('candidate', [''])[0]
                view_type = params.get('viewType', ['interview'])[0]
                
                doc = self.generate_docx(question_ids, candidate_name, view_type)
                
                output = io.BytesIO()
                doc.save(output)
                output.seek(0)
                
                self.send_response(200)
                self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                filename = f"Interview_Questions_{candidate_name.replace(' ', '_')}.docx"
                self.send_header('Content-Disposition', f'attachment; filename="{filename}"')
                self.send_header('Content-Length', str(len(output.getvalue())))
                self.send_header('Access-Control-Allow-Origin', '*')
                self.end_headers()
                self.wfile.write(output.getvalue())
                
            except Exception as e:
                import traceback
                self.send_response(500)
                self.send_header('Content-Type', 'text/plain')
                self.send_header('Access-Control-Allow-Origin', '*')
                self.end_headers()
                self.wfile.write((str(e) + '\n' + traceback.format_exc()).encode())
        else:
            self.send_response(404)
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
    
    def generate_docx(self, question_ids, candidate_name, view_type):
        questions = []
        
        with open('questions.json', 'r') as f:
            all_questions = json.load(f)
        
        for q in all_questions:
            if q['id'] in question_ids:
                questions.append(q)
        
        doc = Document()
        
        # Header with candidate name
        title = doc.add_heading('Interview Questions', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph(f'Candidate: {candidate_name}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f'Total Questions: {len(questions)}')
        doc.add_paragraph('')
        
        if view_type == 'candidate':
            # CANDIDATE VIEW - Only questions
            doc.add_heading('Questions for Candidate', level=1)
            doc.add_paragraph('Solve these questions. No answers provided.')
            doc.add_paragraph('')
            
            for i, q in enumerate(questions, 1):
                p = doc.add_paragraph()
                p.add_run(f'{i}. ').bold = True
                p.add_run(q['question'])
                doc.add_paragraph('')
            
            # Add master copy section
            doc.add_page_break()
            doc.add_heading('Master Copy - For Reference', level=1)
            doc.add_paragraph(f'Candidate: {candidate_name}')
            doc.add_paragraph('')
            
            for i, q in enumerate(questions, 1):
                doc.add_heading(f'{i}. {q["topic"]} - {q["difficulty"]}', level=2)
                
                p = doc.add_paragraph()
                p.add_run('Question: ').bold = True
                p.add_run(q['question'])
                
                p = doc.add_paragraph()
                p.add_run('Answer: ').bold = True
                p.add_run(q['answer'])
                
                doc.add_paragraph('')
        
        else:
            # INTERVIEW VIEW - Full format
            doc.add_heading('Interview Questions - Full Details', level=1)
            doc.add_paragraph('')
            
            for i, q in enumerate(questions, 1):
                doc.add_heading(f'{i}. {q["topic"]} - {q["difficulty"]}', level=2)
                
                p = doc.add_paragraph()
                p.add_run('Topic: ').bold = True
                p.add_run(q['topic'])
                
                p = doc.add_paragraph()
                p.add_run('Difficulty: ').bold = True
                p.add_run(q['difficulty'])
                
                p = doc.add_paragraph()
                p.add_run('Question: ').bold = True
                p.add_run(q['question'])
                
                p = doc.add_paragraph()
                p.add_run('Answer: ').bold = True
                p.add_run(q['answer'])
                
                doc.add_paragraph('')
        
        return doc

print(f"Starting server at http://localhost:{PORT}")
print("Open http://localhost:8000 in your browser")
print("Select questions and click 'Download DOCX' to export")
webbrowser.open(f"http://localhost:{PORT}")

with socketserver.TCPServer(("", PORT), CustomHandler) as httpd:
    httpd.serve_forever()