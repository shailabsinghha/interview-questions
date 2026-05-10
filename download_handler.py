#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import os
import cgi

def generate_docx(question_ids):
    questions = []
    
    with open('questions.json', 'r') as f:
        all_questions = json.load(f)
    
    for q in all_questions:
        if q['id'] in question_ids:
            questions.append(q)
    
    doc = Document()
    
    title = doc.add_heading('Interview Questions', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph(f'Total Selected: {len(questions)} questions')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')
    
    for i, q in enumerate(questions, 1):
        doc.add_heading(f'{i}. {q["topic"]}', level=2)
        
        p = doc.add_paragraph()
        p.add_run('Difficulty: ').bold = True
        p.add_run(q['difficulty'])
        
        p = doc.add_paragraph()
        p.add_run('Question: ').bold = True
        p.add_run(q['question'])
        
        p = doc.add_paragraph()
        p.add_run('Answer: ').bold = True
        p.add_run(q['answer'])
        
        doc.add_paragraph('')
    
    return doc

def handle_download(environ, start_response):
    if environ['REQUEST_METHOD'] == 'POST':
        content_length = int(environ.get('CONTENT_LENGTH', 0))
        post_data = environ['wsgi.input'].read(content_length).decode('utf-8')
        
        try:
            import urllib.parse
            params = urllib.parse.parse_qs(post_data)
            question_ids_str = params.get('questions', [''])[0]
            question_ids = [int(x) for x in question_ids_str.split(',') if x]
            
            doc = generate_docx(question_ids)
            
            response_headers = [
                ('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
                ('Content-Disposition', 'attachment; filename="Interview_Questions.docx"')
            ]
            
            start_response('200 OK', response_headers)
            
            import io
            output = io.BytesIO()
            doc.save(output)
            return [output.getvalue()]
            
        except Exception as e:
            start_response('500 Internal Server Error', [('Content-Type', 'text/plain')])
            return [str(e).encode()]
    else:
        start_response('405 Method Not Allowed', [('Content-Type', 'text/plain')])
        return [b'Method not allowed']

if __name__ == '__main__':
    test_ids = [1, 2, 3, 4, 5]
    doc = generate_docx(test_ids)
    doc.save('test_output.docx')
    print('Test DOCX created')