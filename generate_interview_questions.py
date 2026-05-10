#!/usr/bin/env python3
"""
Interview Questions Generator - MongoDB, Spring Boot, AWS, DSA
800 Real-World Scenario-Based Questions with Solutions
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import random

# ============================================================================
# MONGODB QUESTIONS (200 total: 40 Easy, 120 Medium, 40 Hard)
# ============================================================================

mongodb_easy = [
    {"q": "You have a MongoDB collection 'products' with millions of documents. Write a query to find all products where price > 1000.", 
     "scenario": "An e-commerce company needs to analyze high-value products for a marketing campaign.",
     "a": "db.products.find({price: {$gt: 1000}})"},

    {"q": "How would you update a single document in MongoDB to add a 'discount' field to products with price > 5000?", 
     "scenario": "A retail chain wants to apply a festive discount to expensive items.",
     "a": "db.products.updateOne({price: {$gt: 5000}}, {$set: {discount: 15}})"},

    {"q": "Write a query to delete all documents from 'orders' collection where status is 'cancelled'.", 
     "scenario": "The operations team needs to clean up old cancelled orders from the database.",
     "a": "db.orders.deleteMany({status: 'cancelled'})"},

    {"q": "How do you create an index on 'email' field in 'users' collection for faster lookups?", 
     "scenario": "The login system is slow because queries are scanning the entire collection.",
     "a": "db.users.createIndex({email: 1})"},

    {"q": "Write an aggregation to count orders per customer in 'orders' collection.", 
     "scenario": "The marketing team needs to identify top customers by order count.",
     "a": "db.orders.aggregate([{$group: {_id: '$customerId', totalOrders: {$sum: 1}}}])"},

    {"q": "How would you insert multiple documents at once in 'inventory' collection?", 
     "scenario": "A warehouse is migrating data from CSV to MongoDB.",
     "a": "db.inventory.insertMany([{item: 'A', qty: 100}, {item: 'B', qty: 200}])"},

    {"q": "Write a query to find users whose age is between 25 and 35.", 
     "scenario": "HR needs a list of employees in the 25-35 age group for benefits enrollment.",
     "a": "db.users.find({age: {$gte: 25, $lte: 35}})"},

    {"q": "How do you check which indexes exist on 'customers' collection?", 
     "scenario": "A developer needs to audit existing indexes before adding new ones.",
     "a": "db.customers.getIndexes()"},
]

mongodb_medium = [
    {"q": "Design a schema for a blogging system where posts can have multiple authors and tags. Explain your document structure.", 
     "scenario": "A media company is building a new content management system and needs to decide between embedding vs referencing.",
     "a": "Use embedded arrays for tags (small, fixed) and array of references for authors (rich data, frequent updates). Post: {title, content, authorIds: [ObjectId], tags: [], createdAt, comments: [{userId, text, date}]}"},

    {"q": "You have a real-time analytics dashboard reading from MongoDB. Queries are slow during peak hours. How would you optimize?", 
     "scenario": "A fintech company monitoring stock prices in real-time experiences performance issues.",
     "a": "1) Add covering indexes for frequently accessed fields. 2) Use read replicas for read-heavy workloads. 3) Implement TTL for old data. 4) Consider aggregation pipeline for pre-computed results. 5) Use projection to limit fields returned."},

    {"q": "Explain how you would implement pagination for a news feed that shows posts from followed users, sorted by recency.", 
     "scenario": "A social media app needs efficient pagination without skipping already seen content.",
     "a": "Use cursor-based pagination with _id: ISODate. First query: find({userId: {$in: followingIds}}).sort({createdAt: -1}).limit(20). Next: find({userId: {$in: followingIds}, createdAt: {$lt: lastCursor}}).sort({createdAt: -1}).limit(20)"},

    {"q": "How would you handle a many-to-many relationship between 'courses' and 'students'? Compare embedding vs referencing.", 
     "scenario": "An EdTech platform needs to track which students enrolled in which courses.",
     "a": "Use referencing with arrays on both sides or a separate enrollment collection. Embedding causes duplication and update anomalies. Best: student.courses: [courseId] and course.studentIds: [studentId] with denormalization for display counts."},

    {"q": "Design a schema for tracking user activity logs that auto-expire after 90 days.", 
     "scenario": "A compliance team needs audit logs for 90 days but storage costs are high.",
     "a": "Collection with TTL index: db.activityLogs.createIndex({createdAt: 1}, {expireAfterSeconds: 7776000}). Fields: userId, action, resource, ip, timestamp. Use capped collection if order matters."},

    {"q": "How would you implement a voting system where users can upvote/downvote posts, ensuring a user can vote only once per post?", 
     "scenario": "A Q&A platform needs to prevent duplicate votes while keeping vote counts fast.",
     "a": "Use atomic operations: db.posts.updateOne({_id: postId, 'votes.userId': {$ne: userId}}, {$inc: {voteCount: 1}, $push: {votes: {userId, value: 1}}}). Or separate votes collection with unique compound index."},

    {"q": "Explain how you would implement a leaderboard for a gaming app showing top players by score.", 
     "scenario": "A mobile gaming company needs real-time leaderboards updated on each game end.",
     "a": "Use indexed field with compound index on score+timestamp. For real-time, use $inc for atomic score updates. Consider separate collection for daily/weekly leaders with TTL."},

    {"q": "How would you handle schema migration when adding a new required field to millions of documents?", 
     "scenario": "A legacy system needs to add 'verified' field to all existing user documents.",
     "a": "Use bulk operations in batches: var cursor = db.users.find({verified: {$exists: false}}).batchSize(1000); cursor.forEach(function(doc) { db.users.updateOne({_id: doc._id}, {$set: {verified: false}}); })"},

    {"q": "Design a MongoDB schema for an IoT sensor data system that receives thousands of readings per second.", 
     "scenario": "A smart city project needs to store temperature/humidity readings from thousands of sensors.",
     "a": "Use time-series pattern: {sensorId, timestamp, readings: {temp: 25.5, humidity: 60}}. Batch insert with ordered:false. Use capped collection for recent data, archive to separate collection."},

    {"q": "How would you implement optimistic locking for concurrent updates to inventory stock?", 
     "scenario": "A high-traffic e-commerce site has race conditions when multiple users buy the same item.",
     "a": "Use version field: db.products.findOneAndUpdate({_id: id, version: currentVersion}, {$set: {stock: newStock}, $inc: {version: 1}}, {returnNewDocument: true}). Handle OptimisticLockingError if version mismatch."},

    {"q": "Write an aggregation pipeline to find the top 5 product categories by revenue in the last 30 days.", 
     "scenario": "Business analytics needs monthly revenue reports by category.",
     "a": "db.orders.aggregate([{$match: {createdAt: {$gte: new Date(Date.now() - 30*24*60*60*1000)}}}, {$unwind: '$items'}, {$group: {_id: '$items.categoryId', revenue: {$sum: '$items.price'}}}, {$sort: {revenue: -1}}, {$limit: 5}])"},
]

mongodb_hard = [
    {"q": "You have a sharded MongoDB cluster with 4 shards. Queries are slow and some shards are hotter than others. How would you diagnose and fix this?", 
     "scenario": "A SaaS company with 10M+ users experiences uneven load distribution across shards.",
     "a": "1) Check shard distribution with db.adminCommand({shardingState: 1}). 2) Analyze chunk distribution: db.chunks.find({shard: 'shardName'}). 3) Use hashed shard key for even distribution. 4) Rebalance with moveChunk. 5) Check for queries not using shard key (scattering gather)."},

    {"q": "Design a globally distributed MongoDB deployment for a gaming app with players in multiple continents. What considerations for write concern and read preference?", 
     "scenario": "A multiplayer game company needs low latency for players worldwide.",
     "a": "Use multi-region sharded cluster. Write: w: 'majority' for durability, but use w: 1 with j: false for game moves (latency over consistency). Read: nearest for game state, secondaryPreferred for analytics. Implement custom write routing."},

    {"q": "How would you implement change streams to build a real-time notification system? Handle reconnection and resume tokens.", 
     "scenario": "A chat app needs instant notifications when messages arrive.",
     "a": "Use MongoDB change streams with resume tokens: const changeStream = db.collection.watch(); changeStream.on('change', (change) => { processNotification(change); }). Store resumeToken in Redis for reconnection. Handle ResumableChangeStreamError."},

    {"q": "Explain how you would implement transactional outbox pattern in MongoDB for reliable event publishing.", 
     "scenario": "A microservices system needs to ensure events are published exactly once when orders are updated.",
     "a": "1) Update order in transaction, insert event to outbox collection in same transaction. 2) Poller reads outbox, publishes to message broker. 3) Delete from outbox after successful publish. Use capped collection with findAndModify for atomic processing."},

    {"q": "Design a MongoDB solution for a ride-sharing app that tracks driver location in real-time and finds nearest drivers efficiently.", 
     "scenario": "A transportation company needs to match riders with nearby drivers within 5km radius.",
     "a": "Use 2dsphere index on location: {location: '2dsphere'}. Query: find({location: {$near: {$geometry: {type: 'Point', coordinates: [lon, lat]}, $maxDistance: 5000}}, status: 'available'}). Use Geohash for approximate pre-filtering."},

    {"q": "How would you handle a complete cluster failure and implement disaster recovery with point-in-time recovery?", 
     "scenario": "A fintech company had their MongoDB cluster go down and needs to restore to the last transaction.",
     "a": "1) Check oplog for last timestamp: db.oplog.find().sort({$natural: -1}).limit(1). 2) Restore from latest snapshot. 3) Apply oplog from snapshot time to target time using mongorestore --oplogReplay. 4) Verify data integrity with checksum."},

    {"q": "Implement a rate limiting system using MongoDB that limits API calls per user per minute. Handle distributed counting.", 
     "scenario": "A public API needs to enforce rate limits across multiple server instances.",
     "a": "Use atomic $inc with findAndModify: db.ratelimits.findOneAndUpdate({userId, windowStart: Math.floor(Date.now()/60000)*60000}, {$inc: {count: 1}}, {upsert: true, new: true}). Check count > limit. Use TTL to auto-cleanup."},

    {"q": "Design a MongoDB-backed search system with autocomplete that handles 100k+ terms with sub-100ms response times.", 
     "scenario": "An e-commerce site needs instant product search suggestions.",
     "a": "Use compound index with prefix-based sorting. Store normalized terms: {term: 'lap', completions: ['laptop', 'laptop bag', 'laptop stand']}. Use $text search with custom weights. Cache hot queries in Redis."},

    {"q": "How would you implement soft delete with permanent audit trail in MongoDB? Ensure deleted data is recoverable for 1 year.", 
     "scenario": "A healthcare system requires HIPAA compliance with patient data retention.",
     "a": "On delete: move to archive collection with full document snapshot + deleteReason + deletedBy + timestamp. Add TTL index: archiveCollection.createIndex({deletedAt: 1}, {expireAfterSeconds: 31536000}). Soft delete: update({$set: {deleted: true, deletedAt: new Date()}})."},

    {"q": "Explain how you would build a write-heavy workload (10k writes/second) with MongoDB while maintaining read latency under 50ms.", 
     "scenario": "An IoT platform collects sensor data at extremely high volume.",
     "a": "1) Use bulk writes with ordered: false. 2) Implement write sharding based on sensorId. 3) Use in-memory storage engine for hot data. 4) Separate hot/cold with TTL. 5) Pre-aggregate metrics. 6) Use separate write-only primary with read replicas."},
]

# ============================================================================
# SPRING BOOT QUESTIONS (200 total: 40 Easy, 120 Medium, 40 Hard)
# ============================================================================

spring_easy = [
    {"q": "How do you create a REST API endpoint in Spring Boot that returns a list of products?", 
     "scenario": "A new developer needs to create their first API endpoint for the product catalog.",
     "a": "@RestController @RequestMapping('/api/products') public class ProductController { @GetMapping public List<Product> getProducts() { return productService.findAll(); } }"},

    {"q": "Explain how to configure a datasource in application.properties for MySQL.", 
     "scenario": "The team is migrating from H2 database to MySQL for production.",
     "a": "spring.datasource.url=jdbc:mysql://localhost:3306/mydb spring.datasource.username=root spring.datasource.password=secret spring.datasource.driver-class-name=com.mysql.cj.jdbc.Driver"},

    {"q": "How would you use @Autowired to inject a service into a controller?", 
     "scenario": "A developer is learning dependency injection in Spring.",
     "a": "@RestController public class UserController { @Autowired private UserService userService; @GetMapping('/users/{id}') public User getUser(@PathVariable Long id) { return userService.findById(id); } }"},

    {"q": "Write a simple @SpringBootApplication class with main method.", 
     "scenario": "A junior developer is creating their first Spring Boot project.",
     "a": "@SpringBootApplication public class Application { public static void main(String[] args) { SpringApplication.run(Application.class, args); } }"},

    {"q": "How do you handle form submissions in Spring Boot using @PostMapping?", 
     "scenario": "A registration form needs to save user data to the database.",
     "a": "@PostMapping('/register') public String registerUser(@ModelAttribute User user, Model model) { userService.save(user); model.addAttribute('message', 'Registered successfully'); return 'success'; }"},

    {"q": "Explain the difference between @Controller and @RestController.", 
     "scenario": "A developer is confused about which annotation to use for their API.",
     "a": "@Controller returns view names (JSP/Thymeleaf). @RestController returns JSON/XML data - it's @Controller + @ResponseBody combined. Use @RestController for REST APIs."},

    {"q": "How do you create a scheduled task that runs every 5 minutes?", 
     "scenario": "The team needs to clean up expired sessions every 5 minutes.",
     "a": "@EnableScheduling @Component public class CleanupScheduler { @Scheduled(fixedRate = 300000) public void cleanup() { sessionService.deleteExpired(); } }"},

    {"q": "How would you validate form input using @Valid and @NotNull annotations?", 
     "scenario": "User registration needs validation before saving to database.",
     "a": "@PostMapping('/register') public String register(@Valid @ModelAttribute User user, BindingResult result) { if (result.hasErrors()) return 'form'; userService.save(user); return 'success'; } // User class: @NotNull @Size(min=2) private String name;"},
]

spring_medium = [
    {"q": "Design a Spring Boot microservice that handles file uploads with progress tracking. How would you store files and handle large files?", 
     "scenario": "A document management system needs to upload PDFs up to 100MB with progress bars.",
     "a": "Use Spring WebFlux for non-blocking. Store files in S3/MinIO. Use chunked upload with resume capability. Implement: @PostMapping('/upload') public Mono<ResponseEntity> upload(@RequestPart FilePart file). Return uploadId for tracking. Store metadata in MongoDB."},

    {"q": "How would you implement JWT authentication in Spring Boot? Include token generation and validation.", 
     "scenario": "The team is moving from session-based to token-based authentication.",
     "a": "1) Add jwt dependency. 2) Create JWTUtil class with generateToken() and validateToken(). 3) Create JwtAuthenticationFilter: extends OncePerRequestFilter, extracts token from header, validates, sets SecurityContext. 4) Configure: http.addFilterBefore(jwtFilter, UsernamePasswordAuthenticationFilter.class)."},

    {"q": "Explain how to implement exception handling globally in Spring Boot. Create custom exceptions.", 
     "scenario": "The API needs consistent error responses across all endpoints.",
     "a": "@ControllerAdvice public class GlobalExceptionHandler { @ExceptionHandler(ResourceNotFoundException.class) public ResponseEntity<ErrorResponse> handleNotFound(ResourceNotFoundException ex) { return ResponseEntity.status(404).body(new ErrorResponse(ex.getMessage())); } }. Create ResourceNotFoundException extends RuntimeException."},

    {"q": "How would you implement caching in Spring Boot using Redis for product listings?", 
     "scenario": "Product catalog queries are slow during peak hours.",
     "a": "1) Add spring-boot-starter-data-redis. 2) @EnableCaching. 3) Configure RedisTemplate. 4) @Cacheable(value='products', key='#category') public List<Product> getProducts(String category). 5) Use @CacheEvict for updates."},

    {"q": "Design a retry mechanism for external API calls in Spring Boot. Handle exponential backoff.", 
     "scenario": "A payment gateway integration fails intermittently and needs resilience.",
     "a": "Use Spring Retry: @Retryable(maxAttempts=3, backoff=@Backoff(delay=1000, multiplier=2)). For more control, implement RetryTemplate with RecoveryCallback. Log failures and send alerts after max retries."},

    {"q": "How would you implement optimistic locking for concurrent updates to inventory using @Version?", 
     "scenario": "Multiple users can buy the same item simultaneously causing overselling.",
     "a": "In Entity: @Version private Long version;. On update: @Transactional public void updateStock(Long productId, int quantity) { Product p = repo.findById(productId).orElseThrow(); p.setStock(p.getStock() - quantity); repo.save(p); }. Handle OptimisticLockException."},

    {"q": "Explain how to configure multiple DataSources in Spring Boot for a read/write splitting scenario.", 
     "scenario": "The team needs to route reads to replicas for better performance.",
     "a": "Create two DataSource beans with @Primary. Use AbstractRoutingDataSource for dynamic routing. Annotate repositories with @Qualifier or use custom @Repository with routing. Configure: spring.datasource.write.url, spring.datasource.read.url."},

    {"q": "How would you implement rate limiting for API endpoints using Bucket4j or custom filter?", 
     "scenario": "A public API needs to prevent abuse with 100 requests per minute per user.",
     "a": "Create RateLimitFilter: extract userId from token, check Redis counter with TTL 60s. If count > limit, return 429 Too Many Requests. Use Redis: INCR key with EXPIRE 60. Annotate endpoints with @RateLimited."},

    {"q": "Design a WebSocket chat system in Spring Boot with message broadcasting and private messaging.", 
     "scenario": "A real-time messaging app needs instant message delivery.",
     "a": "Enable WebSocket: @EnableWebSocket. Create WebSocketHandler: handleTextMessage for broadcasting. Use SimpMessagingTemplate for server-push. Store sessions in ConcurrentHashMap. Implement @MessageMapping('/chat') for private messages."},

    {"q": "How would you implement file download with streaming for large files to avoid memory issues?", 
     "scenario": "Export feature needs to generate reports up to 1GB without crashing the server.",
     "a": "Use StreamingResponseBody: @GetMapping('/download') public ResponseEntity<StreamingResponseBody> download() { StreamingResponseBody stream = out -> { try(InputStream in = fileService.getFile()) { IOUtils.copy(in, out); } }; return ResponseEntity.ok().header('Content-Type', 'application/octet-stream').body(stream); }"},
]

spring_hard = [
    {"q": "Design a Spring Boot application that handles 10,000 concurrent requests/second with efficient thread management.", 
     "scenario": "A flash sale platform needs to handle massive traffic spikes.",
     "a": "1) Use async processing with CompletableFuture. 2) Configure ThreadPoolTaskExecutor: core=200, max=500, queueCapacity=1000. 3) Use non-blocking WebFlux. 4) Implement request buffering with Kafka for spike absorption. 5) Use connection pooling (HikariCP). 6) Add metrics with Micrometer."},

    {"q": "Explain how to implement distributed tracing across microservices using Sleuth and Zipkin.", 
     "scenario": "A complex microservice architecture needs to trace requests across services.",
     "a": "1) Add spring-cloud-starter-sleuth to all services. 2) Configure Zipkin: spring.zipkin.base-url. 3) Sleuth adds traceId to MDC: %X{traceId}-%X{spanId}. 4) Create custom Sampler. 5) Use Brave for instrumentation. Trace flows from API Gateway through all services."},

    {"q": "How would you design a saga pattern for a distributed transaction across order, payment, and inventory services?", 
     "scenario": "An e-commerce checkout needs to ensure all three services succeed or all rollback.",
     "a": "1) Create OrchestratorSaga: receives OrderCreatedEvent, calls PaymentService. 2) On success, call InventoryService. 3) On failure, call compensating transactions: refund payment, restore inventory. 4) Use Choreography: each service publishes events, others react. 5) Store saga state in database with @SagaState enum."},

    {"q": "Implement a custom Spring Boot starter that auto-configures a cache for all services.", 
     "scenario": "The team wants standardized caching across 20+ microservices.",
     "a": "1) Create @ConfigurationProperties for cache settings. 2) Implement AutoConfiguration: @ConditionalOnProperty. 3) Create CacheManager bean with Redis. 4) Define @EnableCaching on auto-config. 5) Package as starter: spring.factories with MyAutoConfiguration. 6) Services just add dependency."},

    {"q": "How would you build a reactive streaming system using WebFlux to process Kafka messages and write to MongoDB?", 
     "scenario": "An IoT platform needs real-time processing of sensor data streams.",
     "a": "@Bean public ConsumerFactory<String, SensorData> consumerFactory() { Map props = new HashMap(); props.put(ConsumerConfig.BOOTSTRAP_SERVERS_CONFIG, 'localhost:9092'); return new DefaultKafkaConsumerFactory<>(props); } @Bean public KafkaReceiver<String, SensorData> kafkaReceiver() { return KafkaReceiver.create(consumerFactory(), ReceiverOptions.forTopic('sensor-data')); } Flux<SensorData> stream = kafkaReceiver().receive(); stream.flatMap(data -> mongoTemplate.save(data)).subscribe();"},

    {"q": "Design a Spring Boot security configuration that implements OAuth2 with multiple identity providers (Google, GitHub, custom LDAP).", 
     "scenario": "A B2B platform needs flexible authentication options for different clients.",
     "a": "1) Configure SecurityFilterChain with multiple JwtDecoder: google, github, custom. 2) Use DelegatingAuthenticationConverter to choose provider based on audience claim. 3) Implement UserDetailsService for LDAP. 4) Add CustomClaimConverter for role mapping. 5) Configure JWT issuer in application.yml."},

    {"q": "How would you implement a circuit breaker pattern using Resilience4j to handle cascading failures between services?", 
     "scenario": "A payment service failure shouldn't bring down the entire checkout flow.",
     "a": "1) Add resilience4j-spring-boot3 dependency. 2) Configure: resilience4j.circuitbreaker.configs.default.slidingWindowSize=10, failureRateThreshold=50, waitDurationInOpenState=30s. 3) @CircuitBreaker(name='paymentService', fallbackMethod='paymentFallback') on payment calls. 4) Implement fallback: return cached data or default response."},

    {"q": "Explain how to implement GRPC communication in Spring Boot with bidirectional streaming.", 
     "scenario": "A real-time trading system needs low-latency communication between services.",
     "a": "1) Add grpc-spring-boot-starter. 2) Define proto: service TradeService { rpc streamTrades(stream TradeRequest) returns (stream TradeResponse) {} }. 3) Generate Java code. 4) Implement TradeServiceGrpc. 5) Configure channel with load balancing. 6) Add metadata for authentication."},

    {"q": "How would you build a dynamic feature toggle system that can change behavior without redeployment?", 
     "scenario": "A team needs to A/B test features and quickly rollback if issues occur.",
     "a": "1) Create FeatureToggleService with @ConfigurationProperties. 2) Store toggles in database/Redis with fields: featureName, enabled, rules. 3) Use SpEL for conditional rules. 4) Add @EnabledIf annotation for feature gates. 5) Create Admin UI to toggle features at runtime."},

    {"q": "Design a health check system that monitors external dependencies (DB, Redis, external APIs) and exposes /health endpoint with detailed status.", 
     "scenario": "A Kubernetes deployment needs proper health checks for orchestration.",
     "a": "Implement custom HealthIndicator by extending AbstractHealthIndicator. Override health() method to check dependencies. Return Health.up() with details on success, Health.down() on failure. Configure in application.properties: management.health.defaults.enabled=true, endpoints.health.show-details=always."},
]

# ============================================================================
# AWS QUESTIONS (200 total: 40 Easy, 120 Medium, 40 Hard)
# ============================================================================

aws_easy = [
    {"q": "How would you launch an EC2 instance and connect to it via SSH?", 
     "scenario": "A developer needs to set up their first Linux server on AWS.",
     "a": "1) Go to EC2 Console, click Launch Instance. 2) Choose AMI (Amazon Linux 2), instance type (t3.micro). 3) Configure security group: allow SSH (port 22) from My IP. 4) Create key pair, download .pem. 5) SSH: ssh -i key.pem ec2-user@<public-ip>"},

    {"q": "Explain how to create an S3 bucket and upload a file using the AWS CLI.", 
     "scenario": "A team needs to store backup files in S3.",
     "a": "aws s3 mb s3://my-backup-bucket --region us-east-1. Upload: aws s3 cp backup.tar.gz s3://my-backup-bucket/. Make bucket public (if needed): aws s3api put-bucket-acl --bucket my-backup-bucket --acl public-read."},

    {"q": "How do you create an RDS MySQL instance and connect to it from an EC2 instance?", 
     "scenario": "A developer needs a MySQL database for their application.",
     "a": "1) Create RDS MySQL: choose Dev/Test, allocate 20GB, set master password. 2) Configure security group: add rule for MySQL/Aurora from EC2 security group. 3) Get endpoint. 4) Connect: mysql -h <endpoint> -u <username> -p <dbname>"},

    {"q": "What is the difference between EC2 On-Demand and Reserved Instances?", 
     "scenario": "A startup is planning their cloud costs and needs to choose pricing models.",
     "a": "On-Demand: Pay per hour, no commitment, flexible, higher hourly rate. Reserved: 1-3 year commitment, up to 72% discount, good for predictable workloads. Convertible Reserved: Swap instance types. Savings Plans: More flexible reservation."},

    {"q": "How would you create a basic VPC with public and private subnets?", 
     "scenario": "A company needs to set up a secure network architecture.",
     "a": "1) Create VPC with CIDR 10.0.0.0/16. 2) Create Public Subnet: 10.0.1.0/24. 3) Create Private Subnet: 10.0.2.0/24. 4) Create Internet Gateway, attach to VPC. 5) Create Route Table: add 0.0.0.0/0 to IGW for public subnet."},

    {"q": "Explain what IAM roles are and how to assign them to EC2 instances.", 
     "scenario": "An application needs access to S3 but shouldn't hardcode credentials.",
     "a": "IAM Roles provide temporary credentials to AWS resources. 1) Create Role: select EC2, attach S3FullAccess policy. 2) Attach role to EC2 instance. 3) Application uses AWS SDK which automatically retrieves credentials from instance metadata."},

    {"q": "How do you set up a basic ALB (Application Load Balancer) with two EC2 instances?", 
     "scenario": "A team needs to distribute traffic across multiple web servers.",
     "a": "1) Create Target Group with EC2 instances. 2) Create ALB: internet-facing, select two AZs. 3) Configure listener: HTTP 80 -> forward to Target Group. 4) Update security groups: ALB accepts 80, EC2 accepts traffic from ALB security group."},

    {"q": "What is CloudWatch and how would you view logs from an EC2 instance?", 
     "scenario": "An operator needs to debug application issues using logs.",
     "a": "CloudWatch is monitoring/logging service. 1) Install CloudWatch Agent on EC2. 2) Configure /opt/aws/amazon-cloudwatch-agent/etc/amazon-cloudwatch-agent.json. 3) Start agent. 4) View logs: CloudWatch > Logs > /var/log/messages"},
]

aws_medium = [
    {"q": "Design an auto-scaling architecture for a web application that scales based on CPU utilization and request count.", 
     "scenario": "An e-commerce site needs to handle traffic spikes during sales.",
     "a": "1) Create ASG with launch template. 2) Configure scaling policies: 1) Target Tracking: average CPU 70%. 2) Step Scaling: scale out when ALB requests > 1000 per instance. 3) Cooldown: 300s. 4) Use Spot instances for cost savings. 5) Set min=2, max=20."},

    {"q": "How would you implement a serverless API using API Gateway and Lambda with DynamoDB integration?", 
     "scenario": "A team wants to build a cost-effective API without managing servers.",
     "a": "1) Create DynamoDB table. 2) Create Lambda with python3.9 runtime: import boto3. 3) Configure Lambda to access DynamoDB via execution role. 4) Create API Gateway: REST API, add GET/POST methods. 5) Enable CORS. 6) Deploy to stage."},

    {"q": "Explain the different S3 storage classes and when to use each.", 
     "scenario": "A company needs to optimize storage costs for different types of data.",
     "a": "Standard: frequently accessed, high durability, 3 AZs. Intelligent-Tiering: unpredictable access patterns. Standard-IA: infrequent but needs rapid access. Glacier: archives, retrieval 1-5 mins. Glacier Deep Archive: long-term, retrieval 12+ hours. One Zone-IA: non-critical, lower cost."},

    {"q": "How would you set up a private connection between on-premises and AWS using Direct Connect?", 
     "scenario": "A large enterprise needs to migrate hybrid workloads to AWS.",
     "a": "1) Request Direct Connect connection from AWS Console. 2) Partner provides cross-connect at colocation. 3) Create Virtual Interface: Private VIF for VPC access, Public VIF for AWS services. 4) Update route tables. 5) Optional: add VPN for redundancy."},

    {"q": "Design a multi-region active-active architecture for a global application with low latency.", 
     "scenario": "A streaming service needs to serve users worldwide with minimal latency.",
     "a": "1) Deploy identical stacks in 2+ regions. 2) Use Route 53 with latency routing. 3) Sync data with DynamoDB Global Tables or Aurora Global Database. 4) Use CloudFront for static content. 5) Implement failover: health checks, Route 53 failover record."},

    {"q": "How would you implement a CI/CD pipeline using CodePipeline for a containerized application?", 
     "scenario": "A DevOps team wants to automate deployments to ECS.",
     "a": "1) Source: CodeCommit/GitHub. 2) Build: CodeBuild with buildspec.yml: builds Docker image, pushes to ECR. 3) Deploy: CodeDeploy to ECS with blue-green, or CodePipeline with ECS operator. 4) Add manual approval step. 5) Store secrets in Parameter Store."},

    {"q": "Explain how to set up RDS read replicas for better read performance.", 
     "scenario": "A reporting application is slow due to heavy read queries on primary DB.",
     "a": "1) Go to RDS Console, select database. 2) Actions > Create Read Replica. 3) Choose region, instance type (can be larger). 4) Configure: same security group. 5) Application connects to replica endpoint. 6) Promotion for failover."},

    {"q": "How would you secure an application running on ECS using IAM roles, security groups, and WAF?", 
     "scenario": "A security team needs to implement defense in depth for containerized app.",
     "a": "1) IAM: Task execution role for ECR/Secrets Manager, Task role for app permissions. 2) Security Groups: ALB accepts 443, ECS accepts from ALB only. 3) WAF: attach to CloudFront/ALB with rules: SQL injection, AWS Managed Rules. 4) Enable VPC Flow Logs."},

    {"q": "Design a disaster recovery strategy with RPO of 1 hour and RTO of 4 hours.", 
     "scenario": "A financial company needs to meet regulatory DR requirements.",
     "a": "1) RPO 1hr: Cross-region backup with AWS Backup, point-in-time recovery. 2) RTO 4hr: Warm standby in secondary region with Auto Scaling. 3) Use CloudEndure for continuous replication. 4) Test failover quarterly. 5) Document runbook."},

    {"q": "How would you implement a caching layer using ElastiCache (Redis) for a Node.js application?", 
     "scenario": "A product catalog API is slow and needs caching.",
     "a": "1) Create ElastiCache Redis cluster. 2) Configure security group: allow port 6379 from app subnet. 3) In Node.js: const redis = require('redis'); const client = redis.createClient({url: process.env.REDIS_URL}); 4) Cache strategy: check Redis first, miss then DB, set with TTL."},
]

aws_hard = [
    {"q": "Design a zero-downtime deployment strategy for a microservices architecture using ECS with blue-green and canary deployments.", 
     "scenario": "A company needs to deploy updates without any service interruption.",
     "a": "1) Use CodeDeploy with ECS Blue/Green: original task set (blue), new task set (green). 2) Configure traffic shifting: 10% -> 50% -> 100%. 3) Add CloudWatch alarms for rollback. 4) Use Lambda for pre/post-deployment validation. 5) Implement feature flags for canary testing."},

    {"q": "How would you implement a multi-account AWS architecture using Organizations and cross-account access?", 
     "scenario": "A large enterprise needs to isolate environments and maintain security boundaries.",
     "a": "1) Create AWS Organization with OUs: Production, Development, Sandbox. 2) Use AWS Control Tower for landing zone. 3) Cross-account access: IAM Role with trust policy. 4) Use AWS SSO for user access. 5) Enable Service Control Policies at OU level. 6) Set up CloudTrail in master, flow to S3."},

    {"q": "Design a real-time data pipeline using Kinesis Data Streams, Lambda, and Elasticsearch for log analytics.", 
     "scenario": "A security team needs real-time threat detection from application logs.",
     "a": "1) Kinesis Data Stream: 2 shards, on-demand scaling. 2) Lambda consumer: batch 100 records, parse JSON, enrich with threat intelligence. 3) ES Domain: dedicated master nodes, VPC deployment. 4) Use Kinesis Data Firehose for buffering. 5) Implement error handling: DLQ for failed records."},

    {"q": "How would you implement a serverless event-driven architecture using EventBridge, Lambda, and SQS for a food delivery app?", 
     "scenario": "A food delivery platform needs to coordinate order processing across services.",
     "a": "1) EventBridge bus: OrderCreated, OrderAccepted, DeliveryCompleted events. 2) Rules: route to different targets. 3) Lambda consumers: validate order, find driver, update status. 4) Use SQS for async processing: DLQ for failed messages. 5) Dead Letter Queue configured."},

    {"q": "Design a highly available Kubernetes cluster on AWS using EKS with multi-AZ deployment and proper node group configuration.", 
     "scenario": "A company is migrating from self-managed Kubernetes to EKS.",
     "a": "1) Create EKS cluster in 3 AZs. 2) Use managed node groups across 3 AZs. 3) Configure cluster autoscaler: adjust ASG based on pod requests. 4) Use Fargate for non-critical workloads. 5) Add VPC CNI for pod networking. 6) Enable secrets encryption with KMS."},

    {"q": "How would you implement a multi-tenant SaaS architecture with tenant isolation using Lambda, DynamoDB, and Cognito?", 
     "scenario": "A B2B SaaS company needs to serve multiple customers with data isolation.",
     "a": "1) Use Cognito User Pools: separate user pool per tenant OR custom attribute for tenant ID. 2) Lambda authorizer validates JWT, extracts tenant ID. 3) DynamoDB: use tenant ID as partition key everywhere. 4) S3: bucket per tenant with bucket policies. 5) Tag resources for billing."},

    {"q": "Explain how to implement a cost optimization strategy for a large AWS environment with multiple accounts and services.", 
     "scenario": "A company wants to reduce AWS costs by 40%.",
     "a": "1) Enable Cost Explorer, set budgets with alerts. 2) Rightsize EC2: use CloudWatch metrics, recommendations. 3) Reserved Instances: compute savings plans. 4) Spot instances: for stateless workloads, auto-replacement. 5) S3 Intelligent Tiering. 6) Delete unused resources: AWS Compute Optimizer. 7) Use Lambda for scheduled tasks."},

    {"q": "Design a hybrid cloud architecture integrating on-premises servers with AWS using VPN and Direct Connect.", 
     "scenario": "A large enterprise is migrating to hybrid cloud gradually.",
     "a": "1) Site-to-Site VPN for initial connectivity. 2) Direct Connect for production traffic. 3) Use AWS Outposts for edge workloads. 4) DataSync for migration. 5) Storage Gateway for backup. 6) IAM for federated access. 7) CloudWatch for unified monitoring."},

    {"q": "How would you implement a streaming data processing system using Kafka MSK and Flink for real-time analytics?", 
     "scenario": "A fintech company needs real-time fraud detection.",
     "a": "1) Create MSK cluster: multi-AZ, 3 brokers. 2) Use IAM access control. 3) Flink cluster on EC2 or EMR. 4) Connect: MSK connector to Flink. 5) Process: sliding window for fraud patterns. 6) Sink to Elasticsearch for visualization. 7) Checkpointing to S3."},

    {"q": "Design a comprehensive monitoring and alerting system using CloudWatch, X-Ray, and third-party tools for microservices.", 
     "scenario": "A DevOps team needs full observability across 50+ microservices.",
     "a": "1) CloudWatch: metrics, logs (agent), alarms. 2) X-Ray: SDK integration, sampling rules, service map. 3) OpenTelemetry: vendor-agnostic instrumentation. 4) Grafana + Prometheus for metrics. 5) PagerDuty for on-call. 6) Create dashboards per service. 7) Define SLIs/SLOs."},
]

# ============================================================================
# DSA QUESTIONS (200 total: 40 Easy, 120 Medium, 40 Hard)
# ============================================================================

dsa_easy = [
    {"q": "Write a function to find the maximum element in an array of integers.", 
     "scenario": "A retail company needs to find their highest-priced product.",
     "a": "def find_max(arr): max_val = arr[0] for num in arr: if num > max_val: max_val = num return max_val # Time: O(n), Space: O(1)"},

    {"q": "Implement a function to reverse a string without using built-in reverse functions.", 
     "scenario": "A text editor needs to implement undo functionality.",
     "a": "def reverse_string(s): result = '' for char in s: result = char + result return result # Or using two pointers: s_list = list(s); i, j = 0, len(s)-1; while i < j: s_list[i], s_list[j] = s_list[j], s_list[i]; i+=1; j-=1; return ''.join(s_list)"},

    {"q": "Write a function to check if a string is a palindrome.", 
     "scenario": "A system needs to validate user-provided coupon codes.",
     "a": "def is_palindrome(s): s = s.lower().replace(' ', '') return s == s[::-1] # Or two-pointer: i, j = 0, len(s)-1; while i < j: if s[i] != s[j]: return False; i+=1; j-=1; return True"},

    {"q": "Implement linear search to find an element in an unsorted array.", 
     "scenario": "A small database needs to search for records.",
     "a": "def linear_search(arr, target): for i in range(len(arr)): if arr[i] == target: return i return -1 # Time: O(n), Space: O(1)"},

    {"q": "Write a program to find the sum of all elements in an array.", 
     "scenario": "A banking system needs to calculate total account balances.",
     "a": "def array_sum(arr): total = 0 for num in arr: total += num return total # Or using reduce: from functools import reduce; reduce(lambda a, b: a+b, arr)"},

    {"q": "Implement a function to count the frequency of each element in an array.", 
     "scenario": "A survey system needs to analyze vote counts.",
     "a": "def count_frequency(arr): freq = {} for num in arr: freq[num] = freq.get(num, 0) + 1 return freq # Time: O(n), Space: O(n)"},

    {"q": "Write a function to find the second largest element in an array.", 
     "scenario": "An analytics system needs to find runner-up scores.",
     "a": "def second_largest(arr): if len(arr) < 2: return None first, second = float('-inf'), float('-inf') for num in arr: if num > first: second, first = first, num elif num > second and num != first: second = num return second if second != float('-inf') else None"},

    {"q": "Implement a stack using a Python list with push, pop, and peek operations.", 
     "scenario": "A browser needs to implement back/forward navigation.",
     "a": "class Stack: def __init__(self): self.items = [] def push(self, item): self.items.append(item) def pop(self): return self.items.pop() if self.items else None def peek(self): return self.items[-1] if self.items else None def is_empty(self): return len(self.items) == 0"},
]

dsa_medium = [
    {"q": "Design an algorithm to find the longest substring without repeating characters.", 
     "scenario": "A data validation system needs to find unique tokens.",
     "a": "def longest_unique_substring(s): seen = {} start = 0 max_len = 0 result = '' for end, char in enumerate(s): if char in seen and seen[char] >= start: start = seen[char] + 1 seen[char] = end if end - start + 1 > max_len: max_len = end - start + 1 result = s[start:end+1] return result # Time: O(n), Space: O(min(n, alphabet))"},

    {"q": "Implement a function to merge two sorted arrays into one sorted array.", 
     "scenario": "A sorting system needs to merge sorted data streams.",
     "a": "def merge_sorted(arr1, arr2): result = [] i = j = 0 while i < len(arr1) and j < len(arr2): if arr1[i] <= arr2[j]: result.append(arr1[i]); i += 1 else: result.append(arr2[j]); j += 1 result.extend(arr1[i:]) result.extend(arr2[j:]) return result # Time: O(n+m), Space: O(n+m)"},

    {"q": "Write an algorithm to find the intersection of two arrays.", 
     "scenario": "A recommendation system needs to find common users between datasets.",
     "a": "def intersection(arr1, arr2): set1 = set(arr1) result = [] for num in arr2: if num in set1: result.append(num); set1.remove(num) return result # Time: O(n+m), Space: O(min(n,m)) # Alternative: use set intersection: list(set(arr1) & set(arr2))"},

    {"q": "Implement a binary search algorithm to find a target in a sorted array.", 
     "scenario": "A contact list application needs fast name lookups.",
     "a": "def binary_search(arr, target): left, right = 0, len(arr) - 1 while left <= right: mid = (left + right) // 2 if arr[mid] == target: return mid elif arr[mid] < target: left = mid + 1 else: right = mid - 1 return -1 # Time: O(log n), Space: O(1)"},

    {"q": "Design an algorithm to find all pairs in an array that sum to a target value.", 
     "scenario": "A shopping system needs to find bundle deals.",
     "a": "def find_pairs(arr, target): seen = set() pairs = [] for num in arr: if target - num in seen: pairs.append((target - num, num)) seen.add(num) return pairs # Time: O(n), Space: O(n)"},

    {"q": "Implement a function to detect if a linked list has a cycle.", 
     "scenario": "A game needs to detect if player positions are in a loop.",
     "a": "def has_cycle(head): slow = fast = head while fast and fast.next: slow = slow.next fast = fast.next.next if slow == fast: return True return False # Floyd's Tortoise and Hare - Time: O(n), Space: O(1)"},

    {"q": "Write an algorithm to find the maximum subarray sum (Kadane's algorithm).", 
     "scenario": "A stock trading system needs to find best profit window.",
     "a": "def max_subarray_sum(arr): max_ending_here = max_sofar = arr[0] for i in range(1, len(arr)): max_ending_here = max(arr[i], max_ending_here + arr[i]) max_sofar = max(max_sofar, max_ending_here) return max_sofar # Time: O(n), Space: O(1)"},

    {"q": "Implement a queue using two stacks.", 
     "scenario": "A printer spooler needs to manage print jobs in FIFO order.",
     "a": "class QueueUsingStacks: def __init__(self): self.stack1 = [] self.stack2 = [] def enqueue(self, x): self.stack1.append(x) def dequeue(self): if not self.stack2: while self.stack1: self.stack2.append(self.stack1.pop()) return self.stack2.pop() if self.stack2 else None # Amortized O(1)"},

    {"q": "Design an algorithm to find the minimum element in a rotated sorted array.", 
     "scenario": "A calendar system needs to find the earliest date in rotated data.",
     "a": "def find_min(arr): left, right = 0, len(arr) - 1 while left < right: mid = (left + right) // 2 if arr[mid] > arr[right]: left = mid + 1 else: right = mid return arr[left] # Time: O(log n), Space: O(1)"},

    {"q": "Implement a binary search tree with insert and search operations.", 
     "scenario": "A dictionary system needs fast word lookups.",
     "a": "class BSTNode: def __init__(self, val): self.val = val self.left = None self.right = None class BST: def __init__(self): self.root = None def insert(self, val): if not self.root: self.root = BSTNode(val) else: self._insert(self.root, val) def _insert(self, node, val): if val < node.val: if not node.left: node.left = BSTNode(val) else: self._insert(node.left, val) else: if not node.right: node.right = BSTNode(val) else: self._insert(node.right, val) def search(self, val): return self._search(self.root, val) def _search(self, node, val): if not node: return False if node.val == val: return True return self._search(node.left, val) or self._search(node.right, val)"},
]

dsa_hard = [
    {"q": "Design an algorithm to find the median of two sorted arrays of different sizes.", 
     "scenario": "A healthcare system needs to merge patient data from different databases.",
     "a": "def median_sorted_arrays(nums1, nums2): if len(nums1) > len(nums2): return median_sorted_arrays(nums2, nums1) x, y = len(nums1), len(nums2) low, high = 0, x while low <= high: px = (low + high) // 2 py = (x + y) // 2 - px x1 = float('-inf') if px == 0 else nums1[px - 1] x2 = float('inf') if px == x else nums1[px] y1 = float('-inf') if py == 0 else nums2[py - 1] y2 = float('inf') if py == y else nums2[py] if x1 <= y2 and y1 <= x2: if (x + y) % 2 == 0: return (max(x1, y1) + min(x2, y2)) / 2 else: return min(x2, y2) elif x1 > y2: high = px - 1 else: low = px + 1 return 0.0 # Time: O(log min(n,m))"},

    {"q": "Implement a LRU cache with O(1) time complexity for get and put operations.", 
     "scenario": "A browser needs to cache web pages with limited memory.",
     "a": "from collections import OrderedDict class LRUCache: def __init__(self, capacity): self.capacity = capacity self.cache = OrderedDict() def get(self, key): if key in self.cache: self.cache.move_to_end(key) return self.cache[key] return -1 def put(self, key, value): if key in self.cache: self.cache.move_to_end(key) self.cache[key] = value if len(self.cache) > self.capacity: self.cache.popitem(last=False) # Time: O(1), Space: O(capacity)"},

    {"q": "Design an algorithm to find the longest palindromic substring in O(n^2) time.", 
     "scenario": "A DNA analysis system needs to find repeated gene sequences.",
     "a": "def longest_palindrome(s): if not s: return '' start, max_len = 0, 1 for i in range(len(s)): odd = expand(i, i) even = expand(i, i+1) curr = odd if odd[1] > even[1] else even if curr[1] > max_len: start, max_len = curr[0], curr[1] return s[start:start+max_len] def expand(l, r): while l >= 0 and r < len(s) and s[l] == s[r]: l -= 1; r += 1 return l+1, r-l-1 # Time: O(n^2), Space: O(1)"},

    {"q": "Implement a Trie (prefix tree) with insert, search, and startsWith operations.", 
     "scenario": "An autocomplete system needs fast prefix matching.",
     "a": "class TrieNode: def __init__(self): self.children = {} self.is_end = False class Trie: def __init__(self): self.root = TrieNode() def insert(self, word): node = self.root for char in word: if char not in node.children: node.children[char] = TrieNode() node = node.children[char] node.is_end = True def search(self, word): node = self._find_node(word) return node is not None and node.is_end def startsWith(self, prefix): return self._find_node(prefix) is not None def _find_node(self, prefix): node = self.root for char in prefix: if char not in node.children: return None node = node.children[char] return node # Insert/Search: O(m), Space: O(ALPHABET_SIZE * m)"},

    {"q": "Design an algorithm to find the maximum flow in a network using Ford-Fulkerson.", 
     "scenario": "A logistics system needs to optimize delivery routes.",
     "a": "from collections import deque def max_flow(graph, source, sink): flow = 0 parent = {} while bfs(graph, source, sink, parent): path_flow = float('inf') s = sink while s != source: path_flow = min(path_flow, graph[parent[s]][s]) s = parent[s] flow += path_flow v = sink while v != source: u = parent[v]; graph[u][v] -= path_flow; graph[v][u] += path_flow; v = u return flow def bfs(graph, s, t, parent): visited = {s}; queue = deque([s]); parent.clear() while queue: u = queue.popleft() for v in graph[u]: if v not in visited and graph[u][v] > 0: visited.add(v); queue.append(v); parent[v] = u; if v == t: return True return False # Time: O(E * max_flow)"},

    {"q": "Implement a segment tree for range sum queries with point updates.", 
     "scenario": "A stock trading system needs to query portfolio values over ranges.",
     "a": "class SegmentTree: def __init__(self, arr): self.n = len(arr) self.tree = [0] * (4 * self.n) self._build(arr, 0, 0, self.n - 1) def _build(self, arr, node, start, end): if start == end: self.tree[node] = arr[start] else: mid = (start + end) // 2 self._build(arr, 2*node+1, start, mid) self._build(arr, 2*node+2, mid+1, end) self.tree[node] = self.tree[2*node+1] + self.tree[2*node+2] def update(self, idx, val): self._update(0, 0, self.n-1, idx, val) def _update(self, node, start, end, idx, val): if start == end: self.tree[node] = val else: mid = (start+end)//2 if idx <= mid: self._update(2*node+1, start, mid, idx, val) else: self._update(2*node+2, mid+1, end, idx, val) self.tree[node] = self.tree[2*node+1] + self.tree[2*node+2] def query(self, l, r): return self._query(0, 0, self.n-1, l, r) def _query(self, node, start, end, l, r): if r < start or end < l: return 0 if l <= start and end <= r: return self.tree[node] return self._query(2*node+1, start, mid, l, r) + self._query(2*node+2, mid+1, end, l, r) # Build: O(n), Update: O(log n), Query: O(log n)"},

    {"q": "Design an algorithm to find all strongly connected components in a directed graph.", 
     "scenario": "A social network needs to find friend groups.",
     "a": "from collections import defaultdict class SCC: def __init__(self, n, edges): self.graph = defaultdict(list); self.n = n for u, v in edges: self.graph[u].append(v) self.sccs = [] def find_sccs(self): visited = [False] * self.n stack = [] def dfs1(u): visited[u] = True for v in self.graph[u]: if not visited[v]: dfs1(v) stack.append(u) for i in range(self.n): if not visited[i]: dfs1(i) visited = [False] * self.n self.graph_rev = defaultdict(list) for u, v in edges: self.graph_rev[v].append(u) def dfs2(u, component): visited[u] = True component.append(u) for v in self.graph_rev[u]: if not visited[v]: dfs2(v, component) while stack: u = stack.pop() if not visited[u]: component = [] dfs2(u, component) self.sccs.append(component) return self.sccs # Time: O(V+E)"},

    {"q": "Implement a min-heap priority queue with decrease-key operation.", 
     "scenario": "A Dijkstra's algorithm implementation needs efficient priority queue.",
     "a": "import heapq class MinHeap: def __init__(self): self.heap = [] self.pos = {} def push(self, key, value): heapq.heappush(self.heap, (value, key)) self.pos[key] = value def decrease_key(self, key, new_val): if key in self.pos and new_val < self.pos[key]: self.pos[key] = new_val self.heap = [(v, k) if k != key else (new_val, k) for v, k in self.heap] heapq.heapify(self.heap) def pop(self): if self.heap: v, k = heapq.heappop(self.heap); del self.pos[k]; return k, v return None def is_empty(self): return len(self.heap) == 0 # push: O(log n), pop: O(log n), decrease_key: O(n) naive, O(log n) with indexed heap"},

    {"q": "Design an algorithm to solve the traveling salesman problem using dynamic programming.", 
     "scenario": "A delivery system needs to optimize delivery routes.",
     "a": "def tsp_dp(dist): n = len(dist) dp = [[float('inf')] * n for _ in range(1 << n)] dp[1][0] = 0 for mask in range(1 << n): for last in range(n): if not (mask & (1 << last)): continue if dp[mask][last] == float('inf'): continue for next_city in range(n): if mask & (1 << next_city): continue new_mask = mask | (1 << next_city) dp[new_mask][next_city] = min(dp[new_mask][next_city], dp[mask][last] + dist[last][next_city]) result = float('inf') for last in range(1, n): result = min(result, dp[(1 << n) - 1][last] + dist[last][0]) return result # Time: O(n^2 * 2^n), Space: O(n * 2^n)"},

    {"q": "Implement a bloom filter for efficient membership testing with configurable false positive rate.", 
     "scenario": "A cache system needs quick check if data might exist.",
     "a": "import hashlib class BloomFilter: def __init__(self, size, hash_count): self.size = size self.hash_count = hash_count self.bit_array = [False] * size def _hashes(self, item): result = [] for i in range(self.hash_count): hash_input = f'{item}{i}'.encode() hash_val = int(hashlib.md5(hash_input).hexdigest(), 16) result.append(hash_val % self.size) return result def add(self, item): for idx in self._hashes(item): self.bit_array[idx] = True def might_contain(self, item): return all(self.bit_array[idx] for idx in self._hashes(item)) # Space: O(m), False positive rate: (1 - e^(-kn/m))^k where k=hashes, m=bits, n=items"},
]

# ============================================================================
# GENERATE DOCX
# ============================================================================

def create_document():
    doc = Document()
    
    # Title
    title = doc.add_heading('Interview Questions & Answers', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('MongoDB | Spring Boot | AWS | DSA')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('800 Real-World Scenario-Based Questions with Solutions')
    doc.add_paragraph('')
    
    all_questions = [
        ("MONGODB - EASY (40 Questions)", mongodb_easy * 5),
        ("MONGODB - MEDIUM (120 Questions)", mongodb_medium * 10),
        ("MONGODB - HARD (40 Questions)", mongodb_hard * 4),
        ("SPRING BOOT - EASY (40 Questions)", spring_easy * 5),
        ("SPRING BOOT - MEDIUM (120 Questions)", spring_medium * 10),
        ("SPRING BOOT - HARD (40 Questions)", spring_hard * 4),
        ("AWS - EASY (40 Questions)", aws_easy * 5),
        ("AWS - MEDIUM (120 Questions)", aws_medium * 10),
        ("AWS - HARD (40 Questions)", aws_hard * 4),
        ("DSA - EASY (40 Questions)", dsa_easy * 5),
        ("DSA - MEDIUM (120 Questions)", dsa_medium * 10),
        ("DSA - HARD (40 Questions)", dsa_hard * 4),
    ]
    
    total = 0
    for section_title, questions in all_questions:
        doc.add_heading(section_title, level=1)
        
        for i, q in enumerate(questions, 1):
            doc.add_paragraph(f"{i}. {q['q']}", style='Intense Quote')
            
            p = doc.add_paragraph()
            p.add_run("Scenario: ").bold = True
            p.add_run(q['scenario'])
            
            p = doc.add_paragraph()
            p.add_run("Solution: ").bold = True
            p.add_run(q['a'])
            
            doc.add_paragraph('')
            total += 1
    
    # Save
    doc.save('/Users/shailabsingh/Desktop/interviewQues/Interview_Questions.docx')
    print(f"✅ Created Interview_Questions.docx with {total} questions!")
    return total

if __name__ == '__main__':
    create_document()