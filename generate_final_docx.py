#!/usr/bin/env python3
"""
Interview Questions Generator - With Clickable Table of Contents
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_elem = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_elem.append(rPr)
    new_elem.text = text
    hyperlink.append(new_elem)
    paragraph._p.append(hyperlink)
    return hyperlink

def create_document():
    doc = Document()
    
    title = doc.add_heading('Interview Questions & Answers', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('MongoDB | Spring Boot | AWS | DSA | Troubleshooting')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Real-World Scenario-Based Questions with Solutions')
    doc.add_paragraph('')
    
    doc.add_heading('Table of Contents', level=1)
    doc.add_paragraph('(Click on any topic to jump to that section)')
    doc.add_paragraph('')
    
    toc_items = [
        ("MONGODB - EASY LEVEL", "_mongodb_easy"),
        ("MONGODB - MEDIUM LEVEL", "_mongodb_medium"),
        ("MONGODB - HARD LEVEL", "_mongodb_hard"),
        ("SPRING BOOT - EASY LEVEL", "_spring_easy"),
        ("SPRING BOOT - MEDIUM LEVEL", "_spring_medium"),
        ("SPRING BOOT - HARD LEVEL", "_spring_hard"),
        ("AWS - EASY LEVEL", "_aws_easy"),
        ("AWS - MEDIUM LEVEL", "_aws_medium"),
        ("AWS - HARD LEVEL", "_aws_hard"),
        ("DSA - EASY LEVEL", "_dsa_easy"),
        ("DSA - MEDIUM LEVEL", "_dsa_medium"),
        ("DSA - HARD LEVEL", "_dsa_hard"),
        ("TROUBLESHOOTING SCENARIOS", "_troubleshooting"),
    ]
    
    for i, (item, anchor) in enumerate(toc_items, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. ")
        add_hyperlink(p, item, anchor)
        p.paragraph_format.space_after = Pt(6)
    
    doc.add_page_break()
    
    # Questions data
    sections = [
        ("_mongodb_easy", "MONGODB - EASY LEVEL", [
            {"q": "Write a query to find all products where price > 1000", "s": "E-commerce company needs to analyze high-value products", "a": "db.products.find({price: {$gt: 1000}})"},
            {"q": "How to update documents to add a discount field?", "s": "Retail chain applies festive discounts", "a": "db.products.updateOne({price: {$gt: 5000}}, {$set: {discount: 15}})"},
            {"q": "Query to delete cancelled orders", "s": "Operations team cleans up old data", "a": "db.orders.deleteMany({status: 'cancelled'})"},
            {"q": "Create index on email field", "s": "Login system is slow", "a": "db.users.createIndex({email: 1})"},
            {"q": "Aggregation to count orders per customer", "s": "Marketing needs top customer analysis", "a": "db.orders.aggregate([{$group: {_id: '$customerId', total: {$sum: 1}}}])"},
            {"q": "Insert multiple documents at once", "s": "Warehouse migrating from CSV", "a": "db.inventory.insertMany([{item: 'A', qty: 100}, {item: 'B', qty: 200}])"},
            {"q": "Find users aged between 25 and 35", "s": "HR needs employee list for benefits", "a": "db.users.find({age: {$gte: 25, $lte: 35}})"},
            {"q": "Check existing indexes on collection", "s": "Developer auditing indexes", "a": "db.customers.getIndexes()"},
        ]),
        
        ("_mongodb_medium", "MONGODB - MEDIUM LEVEL", [
            {"q": "Design schema for blogging system with multiple authors", "s": "Media company building CMS", "a": "Use embedded arrays for tags, references for authors"},
            {"q": "Optimize slow real-time analytics queries", "s": "Fintech monitoring stock prices", "a": "Add covering indexes, use read replicas, implement TTL for old data"},
            {"q": "Implement efficient pagination for news feed", "s": "Social media app needs cursor-based pagination", "a": "Use _id and timestamp with $lt for cursor pagination"},
            {"q": "Handle many-to-many relationship for courses and students", "s": "EdTech platform tracking enrollments", "a": "Use referencing or separate enrollment collection"},
            {"q": "Design activity logs with 90-day auto-expiry", "s": "Compliance team needs audit logs", "a": "Use TTL index: createIndex({createdAt: 1}, {expireAfterSeconds: 7776000})"},
            {"q": "Implement voting system preventing duplicate votes", "s": "Q&A platform preventing spam", "a": "Use atomic $inc with unique compound index on votes collection"},
            {"q": "Build real-time gaming leaderboard", "s": "Mobile gaming company needs live scores", "a": "Use $inc for atomic updates, compound index on score+timestamp"},
            {"q": "Schema migration adding required field to millions of docs", "s": "Legacy system adding verified field", "a": "Use bulk operations in batches of 1000"},
            {"q": "Design IoT sensor data schema for high-volume ingestion", "s": "Smart city project with thousands of sensors", "a": "Use time-series pattern with batch inserts"},
            {"q": "Implement optimistic locking for inventory updates", "s": "E-commerce preventing overselling", "a": "Use @Version field with findOneAndUpdate"},
        ]),
        
        ("_mongodb_hard", "MONGODB - HARD LEVEL", [
            {"q": "Diagnose uneven load distribution in sharded cluster", "s": "SaaS company with 10M+ users", "a": "Check shard distribution, use hashed shard key, rebalance chunks"},
            {"q": "Design globally distributed MongoDB for multi-region gaming", "s": "Multiplayer game company", "a": "Use multi-region sharded cluster, w:majority for writes, nearest for reads"},
            {"q": "Implement change streams for real-time notifications", "s": "Chat app needing instant notifications", "a": "Use MongoDB change streams with resume tokens stored in Redis"},
            {"q": "Implement transactional outbox pattern", "s": "Microservices needing reliable events", "a": "Update order and insert to outbox in same transaction"},
            {"q": "Design location-based driver matching system", "s": "Ride-sharing app finding nearby drivers", "a": "Use 2dsphere index with $near query"},
        ]),
        
        ("_spring_easy", "SPRING BOOT - EASY LEVEL", [
            {"q": "Create REST API endpoint returning product list", "s": "Developer creating first API", "a": "@RestController @RequestMapping('/api/products') @GetMapping public List<Product> getProducts()"},
            {"q": "Configure MySQL datasource in application.properties", "s": "Team migrating to MySQL", "a": "spring.datasource.url=jdbc:mysql://localhost:3306/mydb"},
            {"q": "Use @Autowired to inject service into controller", "s": "Developer learning DI", "a": "@Autowired private UserService userService;"},
            {"q": "Write @SpringBootApplication main class", "s": "Junior developer starting project", "a": "@SpringBootApplication public static void main(String[] args)"},
            {"q": "Handle form submission with @PostMapping", "s": "Registration form saving user data", "a": "@PostMapping('/register') public String register(@ModelAttribute User user)"},
            {"q": "Difference between @Controller and @RestController", "s": "Developer confused about annotations", "a": "@RestController returns JSON, @Controller returns views"},
            {"q": "Create scheduled task running every 5 minutes", "s": "Team needing session cleanup", "a": "@Scheduled(fixedRate = 300000) public void cleanup()"},
            {"q": "Validate form input with @Valid and @NotNull", "s": "User registration needing validation", "a": "@Valid @ModelAttribute User user, BindingResult result"},
        ]),
        
        ("_spring_medium", "SPRING BOOT - MEDIUM LEVEL", [
            {"q": "Design file upload microservice with progress tracking", "s": "Document management system with large PDFs", "a": "Use Spring WebFlux, store in S3, return uploadId for tracking"},
            {"q": "Implement JWT authentication in Spring Boot", "s": "Team moving to token-based auth", "a": "Create JWTUtil class, JwtAuthenticationFilter extending OncePerRequestFilter"},
            {"q": "Implement global exception handling", "s": "API needs consistent error responses", "a": "@ControllerAdvice with @ExceptionHandler methods"},
            {"q": "Implement Redis caching for product listings", "s": "Slow product catalog queries", "a": "@EnableCaching, @Cacheable with RedisTemplate"},
            {"q": "Design retry mechanism with exponential backoff", "s": "Payment gateway failing intermittently", "a": "@Retryable(maxAttempts=3, backoff=@Backoff(delay=1000, multiplier=2))"},
            {"q": "Implement optimistic locking with @Version", "s": "Preventing overselling in e-commerce", "a": "@Version private Long version in entity"},
            {"q": "Configure multiple DataSources for read/write splitting", "s": "Route reads to replicas", "a": "Create two DataSource beans with AbstractRoutingDataSource"},
            {"q": "Implement rate limiting for API endpoints", "s": "Prevent API abuse", "a": "Use Redis INCR with TTL 60s, return 429 if exceeded"},
            {"q": "Design WebSocket chat system with broadcasting", "s": "Real-time messaging app", "a": "@EnableWebSocket with SimpMessagingTemplate"},
            {"q": "Implement streaming file download for large files", "s": "Export feature generating 1GB reports", "a": "Use StreamingResponseBody with IOUtils.copy"},
        ]),
        
        ("_spring_hard", "SPRING BOOT - HARD LEVEL", [
            {"q": "Design Spring Boot handling 10,000 concurrent requests", "s": "Flash sale platform with traffic spikes", "a": "Use CompletableFuture, ThreadPoolTaskExecutor, WebFlux, Kafka buffering"},
            {"q": "Implement distributed tracing with Sleuth and Zipkin", "s": "Complex microservice architecture", "a": "Add spring-cloud-starter-sleuth, configure Zipkin, traceId in MDC"},
            {"q": "Design saga pattern for distributed transactions", "s": "E-commerce checkout needing all-or-nothing", "a": "Create orchestrator with compensating transactions"},
            {"q": "Create custom Spring Boot starter for caching", "s": "Team wanting standardized caching", "a": "Use @ConfigurationProperties, AutoConfiguration, spring.factories"},
            {"q": "Build reactive streaming with WebFlux and Kafka", "s": "IoT platform processing sensor data", "a": "Use KafkaReceiver with Flux for streaming"},
        ]),
        
        ("_aws_easy", "AWS - EASY LEVEL", [
            {"q": "Launch EC2 instance and connect via SSH", "s": "Developer setting up first server", "a": "Launch Instance, configure security group, download key pair, ssh -i key.pem"},
            {"q": "Create S3 bucket and upload file via CLI", "s": "Team storing backup files", "a": "aws s3 mb s3://bucket-name && aws s3 cp file s3://bucket-name/"},
            {"q": "Create RDS MySQL and connect from EC2", "s": "Developer needing database", "a": "Create RDS, configure security group, mysql -h endpoint -u user -p"},
            {"q": "Difference between On-Demand and Reserved Instances", "s": "Startup planning costs", "a": "On-Demand: flexible, higher cost. Reserved: commit 1-3 years, up to 72% discount"},
            {"q": "Create basic VPC with public and private subnets", "s": "Company setting up secure network", "a": "Create VPC, IGW, public subnet with route to IGW, private subnet"},
            {"q": "Assign IAM role to EC2 instance", "s": "App needs S3 access without hardcoding", "a": "Create Role, attach policy, assign to EC2"},
            {"q": "Set up ALB with two EC2 instances", "s": "Team distributing traffic", "a": "Create Target Group, create ALB, configure listener"},
            {"q": "Use CloudWatch to view EC2 logs", "s": "Operator debugging issues", "a": "Install CloudWatch Agent, configure json, view in CloudWatch Console"},
        ]),
        
        ("_aws_medium", "AWS - MEDIUM LEVEL", [
            {"q": "Design auto-scaling architecture for web app", "s": "E-commerce handling traffic spikes", "a": "Create ASG, configure Target Tracking CPU 70%, Step Scaling"},
            {"q": "Implement serverless API with API Gateway and Lambda", "s": "Team building cost-effective API", "a": "Create Lambda, add API Gateway, configure triggers"},
            {"q": "Explain S3 storage classes and when to use each", "s": "Company optimizing storage costs", "a": "Standard: frequent. IA: infrequent. Glacier: archives. Deep Archive: long-term"},
            {"q": "Set up Direct Connect for hybrid cloud", "s": "Enterprise migrating workloads", "a": "Request connection, create VIF, update route tables"},
            {"q": "Design multi-region active-active architecture", "s": "Streaming service for global users", "a": "Deploy in 2+ regions, Route 53 latency routing, DynamoDB Global Tables"},
            {"q": "Implement CI/CD pipeline for containerized app", "s": "DevOps automating ECS deployments", "a": "CodePipeline with CodeBuild, CodeDeploy blue-green"},
            {"q": "Set up RDS read replicas for read performance", "s": "Reporting app slow on primary", "a": "Create Read Replica, connect to replica endpoint"},
            {"q": "Secure ECS application with IAM, SG, WAF", "s": "Security implementing defense in depth", "a": "Task execution role, security groups, WAF rules"},
            {"q": "Design DR strategy with RPO 1hr RTO 4hr", "s": "Financial company regulatory requirements", "a": "AWS Backup, warm standby, CloudEndure replication"},
            {"q": "Implement ElastiCache Redis for Node.js", "s": "Slow product API needing cache", "a": "Create Redis cluster, use redis client, cache with TTL"},
        ]),
        
        ("_aws_hard", "AWS - HARD LEVEL", [
            {"q": "Design zero-downtime deployment with ECS blue-green", "s": "Company deploying without interruption", "a": "CodeDeploy blue-green, traffic shifting 10-50-100%, CloudWatch alarms"},
            {"q": "Implement multi-account AWS architecture", "s": "Enterprise isolating environments", "a": "AWS Organizations, OUs, SCPs, AWS SSO"},
            {"q": "Design real-time pipeline with Kinesis and Elasticsearch", "s": "Security team needing threat detection", "a": "Kinesis streams, Lambda consumer, ES domain, Firehose buffering"},
            {"q": "Implement event-driven architecture with EventBridge", "s": "Food delivery coordinating services", "a": "EventBridge bus, rules for routing, Lambda consumers, SQS"},
            {"q": "Design highly available EKS cluster", "s": "Company migrating to EKS", "a": "3 AZs, managed node groups, cluster autoscaler, Fargate"},
        ]),
        
        ("_dsa_easy", "DSA - EASY LEVEL", [
            {"q": "Find maximum element in array", "s": "Retail finding highest-priced product", "a": "Iterate and track max. Time: O(n), Space: O(1)"},
            {"q": "Reverse string without built-in functions", "s": "Text editor implementing undo", "a": "Two-pointer or build new string. Time: O(n)"},
            {"q": "Check if string is palindrome", "s": "Validating coupon codes", "a": "Two-pointer from both ends. Time: O(n)"},
            {"q": "Implement linear search", "s": "Small database record search", "a": "Iterate and compare. Time: O(n)"},
            {"q": "Find sum of array elements", "s": "Banking calculating total balances", "a": "Loop and accumulate. Time: O(n)"},
            {"q": "Count frequency of elements", "s": "Survey analyzing votes", "a": "HashMap for counting. Time: O(n)"},
            {"q": "Find second largest element", "s": "Analytics finding runner-up", "a": "Track largest and second. Time: O(n)"},
            {"q": "Implement stack with push/pop/peek", "s": "Browser back/forward navigation", "a": "List-based implementation with O(1) operations"},
        ]),
        
        ("_dsa_medium", "DSA - MEDIUM LEVEL", [
            {"q": "Find longest substring without repeating chars", "s": "Data validation finding unique tokens", "a": "Sliding window with HashMap. Time: O(n), Space: O(min(n, alphabet))"},
            {"q": "Merge two sorted arrays", "s": "Sorting system merging streams", "a": "Two-pointer merge. Time: O(n+m)"},
            {"q": "Find intersection of two arrays", "s": "Recommendation finding common users", "a": "Set-based intersection. Time: O(n+m)"},
            {"q": "Binary search in sorted array", "s": "Contact list fast lookup", "a": "Standard binary search. Time: O(log n)"},
            {"q": "Find pairs summing to target", "s": "Shopping finding bundle deals", "a": "HashSet approach. Time: O(n)"},
            {"q": "Detect cycle in linked list", "s": "Game detecting position loops", "a": "Floyd's Tortoise and Hare. Time: O(n)"},
            {"q": "Maximum subarray sum (Kadane's)", "s": "Stock trading finding profit window", "a": "Kadane's algorithm. Time: O(n)"},
            {"q": "Implement queue using two stacks", "s": "Printer spooler FIFO order", "a": "Amortized O(1) push, O(1) pop"},
            {"q": "Find minimum in rotated sorted array", "s": "Calendar finding earliest date", "a": "Binary search variation. Time: O(log n)"},
            {"q": "Implement BST with insert and search", "s": "Dictionary fast word lookup", "a": "Standard BST operations. Time: O(log n) average"},
        ]),
        
        ("_dsa_hard", "DSA - HARD LEVEL", [
            {"q": "Find median of two sorted arrays", "s": "Healthcare merging patient data", "a": "Binary search on smaller array. Time: O(log min(n,m))"},
            {"q": "Implement LRU cache with O(1) operations", "s": "Browser caching web pages", "a": "OrderedDict with move_to_end. Time: O(1)"},
            {"q": "Find longest palindromic substring", "s": "DNA finding gene sequences", "a": "Expand around center. Time: O(n^2)"},
            {"q": "Implement Trie for autocomplete", "s": "Autocomplete fast prefix matching", "a": "Node-based prefix tree. Insert/Search: O(m)"},
            {"q": "Find maximum flow using Ford-Fulkerson", "s": "Logistics optimizing routes", "a": "BFS for augmenting paths. Time: O(E * max_flow)"},
        ]),
        
        ("_troubleshooting", "TROUBLESHOOTING SCENARIOS", [
            {"q": "MongoDB queries running slow - diagnose", "s": "E-commerce slow product searches", "a": "Use explain(), add indexes, check projections, review collection size"},
            {"q": "Spring Boot port already in use", "s": "Developer app failing to start", "a": "lsof -i :8080, kill process or change port in properties"},
            {"q": "EC2 instance not accessible via SSH", "s": "Admin cannot reach production server", "a": "Check security group, instance status, public IP, key permissions"},
            {"q": "Kubernetes pod crashing with OOMKilled", "s": "Microservices failing in prod", "a": "Check logs, describe pod, increase memory, profile for leaks"},
            {"q": "Lambda function timing out frequently", "s": "Image processing Lambda slow", "a": "Increase timeout, optimize code, use provisioned concurrency"},
            {"q": "RDS database at 100% CPU", "s": "Production database unresponsive", "a": "Check CloudWatch, run SHOW PROCESSLIST, add indexes, use read replicas"},
            {"q": "S3 bucket returning 403 Forbidden", "s": "App cannot access S3 files", "a": "Check bucket policy, IAM role, ACLs, verify region"},
            {"q": "ALB returning 502 Bad Gateway", "s": "Users seeing 502 errors", "a": "Check target health, security groups, app logs, health checks"},
            {"q": "Docker container failing with connection refused", "s": "Container not starting", "a": "Check docker logs, port mappings, dependencies, env variables"},
            {"q": "API returning inconsistent data", "s": "Frontend seeing different data", "a": "Check race conditions, transaction commits, caching, load balancer"},
        ]),
    ]
    
    total = 0
    for anchor, section_title, questions in sections:
        doc.add_page_break()
        
        heading = doc.add_heading(section_title, level=2)
        heading.anchor = anchor
        
        doc.add_paragraph(f"Total Questions: {len(questions)}")
        doc.add_paragraph('')
        
        for i, q in enumerate(questions, 1):
            p = doc.add_paragraph(f"{i}. {q['q']}")
            p.runs[0].bold = True
            
            p = doc.add_paragraph()
            p.add_run("Scenario: ").bold = True
            p.add_run(q['s'])
            
            p = doc.add_paragraph()
            p.add_run("Solution: ").bold = True
            p.add_run(q['a'])
            
            doc.add_paragraph('')
            total += 1
    
    doc.add_page_break()
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(f"Total Questions: {total}")
    doc.add_paragraph('')
    doc.add_paragraph("Topics Covered:")
    doc.add_paragraph("- MongoDB: Queries, indexing, aggregation, sharding, replication")
    doc.add_paragraph("- Spring Boot: REST APIs, security, caching, microservices")
    doc.add_paragraph("- AWS: EC2, S3, RDS, Lambda, VPC, Containers, Serverless")
    doc.add_paragraph("- DSA: Arrays, Strings, Trees, Graphs, Dynamic Programming")
    doc.add_paragraph("- Troubleshooting: Performance, connectivity, debugging")
    
    doc.save('/Users/shailabsingh/Desktop/interviewQues/Interview_Questions.docx')
    print(f"✅ Created Interview_Questions.docx with {total} questions!")
    print(f"🔗 Clickable Table of Contents with {len(sections)} sections")

if __name__ == '__main__':
    create_document()